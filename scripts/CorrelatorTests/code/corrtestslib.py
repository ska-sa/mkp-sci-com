# Now follow functions that are broken off from process_corrtest.py etc

import numpy as N
import mylibs, libs, katdal
import os, sys, time, glob
import corr_test_chans as cc
import corrtests_config as cfg
from importlib import reload
from katsdpcal.calprocs import get_bls_lookup, k_fit, ants_from_bllist, normalise_complex, g_fit

from docx import Document
from docx.shared import Inches

import matplotlib
matplotlib.use("Agg")
import pylab as pl
pl.ion()

def get_cbid(sysargv):
  """
  If no input on command line, take the cbid from the rdb file in pwd
  """

  if len(sysargv)==1:
    fns = glob.glob('15*rdb')
    if len(fns)!=1: raise RuntimeError
    return fns[0].split('_')[0]
  else:
    return sysargv[1]


################################################################################################
def setup_objd(fname, basedir, doc):
  """
  Set up object d, document and create directories
  """

  print("Setting up the object d, with documentation as", doc)
  class myclass(object):
      pass

  d = myclass()
  d.doc = doc
  mylibs.create_dirs(d, fname, basedir)
  if doc: mylibs.dodocs(d)

  return d
  
################################################################################################
def setup_file(d):
  """
  Set up file from kat open, read in parameters, put onto object, get scans
  """
  print("Reading in rdb file, set up parameters for object d, get scan info")
  t1 = time.time()
  fn = glob.glob(d.fdir+d.fname+'*rdb')[0]
  f = katdal.open(fn)
  t2 = time.time()
  print("Done opening", fn, " in ", t2-t1, " secs")
  
  mylibs.read_katdal_para(f, d)
  
  d.ant_names = [a.name for a in f.ants]
  d.cross_blslook = get_bls_lookup(d.ant_names, f.corr_products)[d.na*4:d.na*4+d.nbl]
  d.index0=[cp[0] for cp in d.cross_blslook]
  d.index1=[cp[1] for cp in d.cross_blslook]
  d.index0, index1 = N.asarray(d.index0), N.asarray(d.index1)
  d.antnums = [int(a[1:]) for a in d.ant_names]
  d.npol = 2
  d.corrs = ['auto', 'cross']; d.corrname = ['Ant', 'Baseline']

  # config parameters
  d.ringthresh = 0.67
  defaults = {'doc': True, 'basedir': '/data/mohan/' , 'ring_nmax': 10 , 'ring_win': 31 , 'fft_win': 51 , 'fft_thresh': 10.0 ,\
              'fft_flag': 10 , 'fft_nmax': 15 , 'detrend_win': 51 , 'cc_thresh': 3.0 , 'cc_minlen': 10 , 'comb_mnn': 2 ,\
              'comb_nmax': 100 , 'comb_nsearch': 1 , 'comb_minfold': 5 , 'comb_thresh': 5.0} 
  d.cfg = cfg
  print()
  for dum in list(defaults.keys()):
    if dum not in cfg.__dir__():
      setattr(d.cfg, dum, defaults[dum])
  
  # Print out details
  f.select()
  #print(f)
  if d.doc:
    d.docfull.add_heading("CONTENTS")
    d.docfull.add_paragraph(str(f))
  
  f.select(scans='track')
  d.gscans = []; d.sind = {}
  d.gtimeinds = []
  ii = 0
  for s in f.scans():
    d.gscans.append(s[0])
    d.gtimeinds.append(f.vis.shape[0])
    d.sind[s[0]] = ii
    ii += 1
  d.gtimeinds, d.gscans = N.asarray(d.gtimeinds), N.asarray(d.gscans)
  d.gnscan = len(d.gscans)
  d.gntime = N.sum(d.gtimeinds)
  
  d.tind0, d.dtind = N.zeros(d.gnscan,int), N.zeros(d.gnscan,int)
  for i in range(d.gnscan):
      if i>0: d.tind0[i] = N.sum(d.dtind)
      d.dtind[i] = d.gtimeinds[i]

  print("\nGood scans are")
  sources = []
  f.select()
  for g in d.gscans:
      f.select(scans=g)
      sources.append(f.catalogue.targets[f.target_indices[0]].name)
      print(g, f.vis.shape, sources[-1])
  print("\nNumber of good dumps is ", N.sum(d.gtimeinds),"\n")
  if d.doc:
    d.docsumm.add_heading("Scans used", 2)
    for ii,g in enumerate(d.gscans):
        str1 = "Scan %i, %s %3i dumps,\t %s" %(g, '  ', d.dtind[ii], sources[ii])
        d.docsumm.add_paragraph(str1, style='List Bullet')

################################################################################################

def divide0(aa, x):
    for i in range(len(aa)):
          aa[i] = [int(aa[i][0]/x), int(aa[i][1]/x)]
    return aa

def divide1(aa, x):
    for i in range(len(aa)):
        for j in range(len(aa[i])):
            aa[i][j] = [int(aa[i][j][0]/x), int(aa[i][j][1]/x)]
    return aa


################################################################################################

def setup_bands(d):

  reload(cc)
  print("Setting up bands to use")
  d.bands = divide0(cc.mybands[d.band][32],32/d.nch)
  d.bandflags = divide1(cc.mybandflags[d.band][32],32/d.nch)
  d.nband = cc.mynband[d.band][32]
  d.fullband = [int(cc.myfullband[d.band][32][0]/(32/d.nch)), \
              int(cc.myfullband[d.band][32][1]/(32/d.nch))]
  d.fullbandflags = divide0(cc.myfullbandflags[d.band][32],32/d.nch)
  
  if os.path.isfile('flagchans'):
    d.fullbandflags = []
    for line in open('flagchans'):
      line = N.asarray(line.strip().split(' '), int)
      d.fullbandflags.append([line[0], line[1]])
  print('Band: ',d.fullband,"\n")
  
  if d.doc:
      d.docsumm.add_paragraph("Bands used")
      d.docfull.add_paragraph("Bands used")
      for ii in range(d.nband):
          mylibs.bothdoc(d, 'para', "Band chs "+str(ii)+": "+str(d.bands[ii][0])+'-'+str(d.bands[ii][1]), style='List Bullet')
      if d.fullband != None:
          mylibs.bothdoc(d, 'para', "Full band chs : "+str(d.fullband[0])+'-'+str(d.fullband[1]), style='List Bullet')
  
  
################################################################################################

def sdpflag_summary(d):
  """
  Read the SDP flags from the file (default is not full rdb file) and plot summary stats of
  flagging fraction as a function of time, channel and baseline/antenna for visibility and
  auto corr data.
  """
  # First see how the sdp has flagged autocorr and then crosscorr
  # Nansum (sdpflags 0,1->0,N.nan) over all chans to get flags as pol, time, chan, ant/bl
  # Do sum over everything to get av flag for each of the above
  print("Calculate summary of SDP flagging fraction")
  if d.doc: d.docfull.add_heading("Flagging summary", 2)
  t1 = time.time()
  ncorr = [d.nant, d.nbl]
  flagfrac=N.zeros((2,2))
  if d.doc: d.docsumm.add_heading("Flagging summary", 2)
  for icorr,corr in enumerate(d.corrs):
      print("Processing ",corr,"scan", end=' ')
      fa_chan, fa_corr = [], []
      for ii,iscan in enumerate(d.gscans):
          mylibs.mywrite(iscan)
          mydir = d.sdpdir+'scan_'+str(iscan)+'/'
  
          fl = mylibs.myload(mydir+'scan_'+str(iscan)+'_full'+corr+'flags')
          fa_chan.extend(N.moveaxis(N.nansum(fl,3),0,1))  # pol,time,chan -> time,pol,chan
          fa_corr.extend(N.moveaxis(N.nansum(fl,2),0,1))  # pol,time,corr -> time,pol,corr; 1,N.nan->n,0
  
      print()
      for ipol in range(d.npol):
          flagfrac[icorr,ipol] = N.nansum(N.asarray(fa_chan)[:,ipol,:])*1.0/(d.nchan*N.sum(d.gtimeinds)*ncorr[icorr])
  
      dumr = N.copy(N.asarray(fa_corr))
      fa_chan, fa_corr = N.asarray(fa_chan), N.asarray(fa_corr)
      fa_time = N.transpose(N.nansum(fa_corr,2))/(ncorr[icorr]*d.nchan)
      fa_chan = N.nansum(fa_chan,0)/(N.sum(d.gtimeinds)*ncorr[icorr])
      fa_corr = N.nansum(fa_corr,0)/(N.sum(d.gtimeinds)*d.nchan)
      fa_time = fa_time[:,1:]  # ignore dc
  
      pl.figure(figsize=(7,10))
      fa_time = fa_time[:,1:-2]
      fa = [fa_chan, fa_corr, fa_time]; titles = ['Channel', 'Corr', 'Time']
      for i in range(3):
          pl.subplot(3,1,i+1)
          pl.plot((1-fa[i][0])*100); pl.plot((1-fa[i][1])*100)
          pl.xlabel(titles[i]); pl.ylabel('SDP flag %');
      pl.suptitle(d.fname+' sdp flag summary for '+corr)
      pl.savefig(d.pldir+d.fname+'_sdp_flag_summary_'+corr+'.png')
  
      for ipol in range(d.npol):
        str1 = "%s %s %s %i %s %.1f %s" %("  Percentage of ",corr,"flags in pol",\
               ipol,'is',(1.0-flagfrac[icorr,ipol])*100.0, "%")
        print(str1)
        if d.doc:
          mylibs.bothdoc(d, 'para', str1, style="List Bullet")
      del fl
  
  if d.doc:
      d.docfull.add_picture(d.pldir+d.fname+'_sdp_flag_summary_'+d.corrs[0]+'.png', \
                           width=Inches(5.0))
      d.docfull.add_picture(d.pldir+d.fname+'_sdp_flag_summary_'+d.corrs[1]+'.png', \
                           width=Inches(5.0))
  t2 = time.time()
  print("%s %.1f %s\n" %("Done in", t2-t1, "sec"))



################################################################################################

def statsbyscan(d):
  """
  Plot mean and rms of the visibility autocorr amps by scan.
  """
  #Plot mean and variance of cross and auto for each scan-av data
  #This is npol X nchan X ncorr X nscan. Do mean and var over chan
  #Mean is direct and var is after poly_filter

  if d.doc: d.docfull.add_heading("Spectral mean and variance", 2)
  t1 = time.time()
  ncorr = [d.nant, d.nbl]
  print("Statistics of scan-averaged data")
  baddata, badtitles = [], []
  for iprod,corr in enumerate(d.corrs):
      f1, f2 = pl.figure(figsize=(9,10)), pl.figure(figsize=(9,10)); ff = [f1,f2]
      stats = N.zeros((2,d.nband,d.npol,d.gnscan,ncorr[iprod]))
      print("Processing", corr, "scans", end=' ')
      for iscan,scan in enumerate(d.gscans):
          data = mylibs.myload(d.sdpdir+'scan_'+str(scan)+'/scan_'+str(scan)+'_'+corr+'_scanav')
          mylibs.mywrite(scan)
          for iband in range(d.nband):
              bchan, echan = d.bands[iband]
              dum = N.abs(data[0][:,bchan:echan,:])  # pol chan corr
              for iflags in range(len(d.bandflags[iband])):
                  dum[:,d.bandflags[iband][iflags][0]-bchan:d.bandflags[iband][iflags][1]-bchan] = N.nan
              stats[0][iband,:,iscan,:] = N.nanmean(dum,1)
              for ipol in range(d.npol):
                  for icorr in range(ncorr[iprod]):
                      x = dum[ipol,:,icorr]
                      stats[1][iband,ipol,iscan,icorr] = N.nanstd(x-mylibs.poly_filter(x,31))*stats[0][iband,ipol,iscan,icorr]
      print()
  
      for iband in range(d.nband):
          for ipol in range(d.npol):
              for ii in range(2):   # mean, rms
                  pl.figure(ff[ii].number)
                  pl.subplot(d.nband,d.npol,iband*d.npol+ipol+1)
                  libs.imshow(N.transpose(stats[ii][iband,ipol]))
                  pl.yticks(N.arange(d.gnscan), d.gscans)
                  if ipol==0: pl.ylabel('Scan')
                  if iband==d.nband-1: pl.xlabel(d.corrname[iprod])
                  pl.colorbar()
                  pl.title('Band'+str(iband)+' pol'+str(ipol))
      pl.figure(ff[0].number)
      pl.suptitle('Mean of ' +d.corrs[iprod]+' for each scan and '+d.corrname[iprod]+' per pol')
      pl.savefig(d.pldir+d.fname+'_'+d.corrs[iprod]+'_mean_scan-ant-pol.png')
      pl.figure(ff[1].number)
      pl.suptitle('RMS of detrended '+d.corrs[iprod]+' for each scan and '+d.corrname[iprod]+' per pol')
      pl.savefig(d.pldir+d.fname+'_'+d.corrs[iprod]+'_rmspf_scan-ant-pol.png')
      if d.doc:
          d.docfull.add_picture(d.pldir+d.fname+'_'+d.corrs[iprod]+'_mean_scan-ant-pol.png', width=Inches(5.0))
          d.docfull.add_picture(d.pldir+d.fname+'_'+d.corrs[iprod]+'_rmspf_scan-ant-pol.png', width=Inches(5.0))
  t2 = time.time()
  print("%s %.1f %s\n" %("Done in", t2-t1, "sec"))


################################################################################################
def plot_ringing(d, baddata, badtitles, nmax, suptitle, savename):

    nn = len(baddata)
    if nn>0:
        print(nn, nmax)
        if nn>nmax: # plot nmax worst
            vals = N.zeros(nn)
            for ii in range(nn):
                vals[ii] = baddata[ii][3]
            indsort = N.argsort(vals)[::-1]
            indsort = indsort[:nmax]
        else:
            indsort = N.arange(nn)

        print(indsort)
        pl.figure()
        for ii,i in enumerate(indsort):
            x = baddata[i][0]; x = x-mylibs.poly_filter(x,31)
            ax = pl.subplot(len(indsort),2,2*ii+1); pl.plot(x); pl.title(badtitles[i])
            if i<len(baddata)-1: ax.set_xticks([])
            pl.subplot(len(indsort),2,2*ii+2); pl.plot(baddata[i][1])
            ind1 = baddata[i][2]
            pl.plot(ind1,baddata[i][1][ind1], '.r')
            if ii==0: pl.title('Ringing index')
        pl.suptitle(suptitle+"  (10 worst)")
        pl.savefig(d.pldir+d.fname+savename)

        d.docfull.add_picture(d.pldir+d.fname+savename,width=Inches(5.0))

################################################################################################

def ringing_2chan(d, thresh=None):
  
  # Now look for 2chan spectral ringing in scan averaged data (auto and corr)
  # sep cross out and pprocess it
  print("Searching for 2-channel ringing")
  t1 = time.time()
  if thresh==None: thresh = d.ringthresh
  baddata, badtitles = [], []
  if d.doc:
      mylibs.bothdoc(d, "title", "2-channel ringing")
      mylibs.bothdoc(d, "para", "Using threshold of ringing index of "+str(thresh) + "\n(1->perfect ringing, 0-> none, neg->higher periods)")
  
  ringinds = [N.zeros((d.npol,d.nant,d.gnscan)),N.zeros((d.npol,d.nbl,d.gnscan))]  # ringing index
  for iprod,corr in enumerate(d.corrs):
      print("Processing", corr, " scan", end=' ')
      bad = 0
      for iscan,scan in enumerate(d.gscans):
          data = mylibs.myload(d.sdpdir+'scan_'+str(scan)+'/scan_'+str(scan)+'_'+corr+'_scanav')
          mylibs.mywrite(scan)
          for iband in range(d.nband):
              bchan, echan = d.bands[iband]
              dum = N.abs(data[0][:,bchan:echan,:])
              for iflags in range(len(d.bandflags[iband])):
                  dum[:,d.bandflags[iband][iflags][0]-bchan:d.bandflags[iband][iflags][1]-bchan] = N.nan
              ret = mylibs.chan2ripple(dum, 1, domf=True, win=d.cfg.ring_win)
              for ipol in range(d.npol):
                  ringinds[iprod][ipol,:,iscan] = ret[ipol]
                  nn = dum.shape[1]
                  ind = N.where(ret[ipol]>=thresh)[0]
                  if len(ind)>0:
                      print("  Band ",iband, " Pol", ipol, ',',len(ind),d.corrname[iprod], 'with 2-chan ringing in', corr, 'thresh',thresh)
                      print("  "+d.corrname[iprod]+" are", ind)
                      print("  Ringing index",ret[ipol][ind])
                      for ind1 in ind:
                          baddata.append([dum[ipol,:,ind1], ret[ipol], ind1, ret[ipol][ind1]])
                          badtitles.append(corr+' S'+str(scan)+' B'+str(iband)+' P'+str(ipol)+' C'+str(ind1)+" %.2f" %(ret[ipol][ind1]))
                      bad += 1
      if bad==0: print("\n%s %.2f %s" %("  No 2-chan ringing found (threshold=",thresh,")"))
      if d.doc: mylibs.bothdoc(d, 'para', "Num of bad "+corr+" is "+str(bad), style="List Bullet")
  print("Done")
  plot_ringing(d, baddata, badtitles, d.cfg.ring_nmax, "2-chan ringing in scanav auto and cross", "_2chanring_scanav.png")
  plot_ringind(d, ringinds, N.arange(d.gnscan), d.gscans, 'spectra')

  t2 = time.time()
  print("%s %.1f %s\n" %("Done in", t2-t1, "sec"))
  
  
################################################################################################

def ringing_2dumpdc(d, thresh=None, mindump=20):
  
  # Now look for 2chan time ringing in dc power for each scan and ant/bl (auto and corr)
  # sep cross out and pprocess it
  t1 = time.time()
  ncorr = [d.nant, d.nbl]
  if thresh==None: thresh=d.ringthresh
  baddata, badtitles = [], []
  print("Searching for 2-dump ringing in dc power using threshhold", thresh)
  mindump = 20
  if d.doc:
      mylibs.bothdoc(d, "title", "2-dump ringing in DC")
      mylibs.bothdoc(d, "para", "Threshold = "+str(thresh)+"\n"+ "Plotting for scans >"+str(mindump)+"dumps)")
  # plot ringing index only for cross for scans with >20 dumps
  indring = N.where(d.gtimeinds>mindump)[0]
  ringinds = [N.zeros((d.npol,d.nant,len(indring))),N.zeros((d.npol,d.nbl,len(indring)))]  # ringing index
  for icorr,corr in enumerate(['auto', 'cross']):
      print("Processing", corr)
      bad = 0; jj = 0
      for iscan,scan in enumerate(d.gscans):
          data = mylibs.myload(d.sdpdir+'scan_'+str(scan)+'/scan_'+str(scan)+'_full'+corr+'dc')
          dum = N.abs(data) # npol X ntime X ncorr
          ret = mylibs.chan2ripple(dum, 1, domf=False)  # npol X ncorr
          if d.gtimeinds[iscan]>mindump:
              ringinds[icorr][:,:,jj] = ret
              jj += 1
          print("For", corr, "scan   ",scan)
          for ipol in range(d.npol):
              nn = dum.shape[1]
              ind = N.where(ret[ipol]>=thresh)[0]
              if len(ind)>0:
                  print("  Pol", ipol, ',',len(ind),d.corrname[icorr], 'with 2-dump ringing in', corr, 'thresh',thresh)
                  for ind1 in ind:
                      baddata.append([dum[ipol,:,ind1], ret[ipol], ind1, ret[ipol][ind1]])
                      badtitles.append(corr+' S'+str(scan)+' P'+str(ipol)+ ' C'+str(ind1)+" %.2f" %(ret[ipol][ind1]))
                  bad += 1
      if d.doc: mylibs.bothdoc(d, 'para', "Num of bad "+corr+" is "+str(bad),style="List Bullet")
      if bad==0: print("No 2-dump ringing in dc found in", corr)
  plot_ringing(d, baddata, badtitles, d.cfg.ring_nmax, "2-dump ringing in DC auto and cross", "_2dumpring_dc.png")
  plot_ringind(d, ringinds, N.arange(jj), d.gscans[indring], 'dc')

  t2 = time.time()
  print("%s %.1f %s\n" %("Done in", t2-t1, "sec"))
  
################################################################################################

def plot_ringind(d, ringinds, xx, yy, title):
  
    for iprod in range(len(d.corrs)):
      pl.figure(figsize=(10,6))
      for ipol in range(d.npol):
          pl.subplot(2,1,ipol+1)
          if N.max(ringinds[iprod][ipol])>d.ringthresh: libs.imshow(ringinds[iprod][ipol],vmin=0.,vmax=1.)
          else: libs.imshow(ringinds[iprod][ipol])
          pl.colorbar()
          pl.yticks(xx, yy)
          pl.title('Pol'+str(ipol)); pl.ylabel('Scan')
          if ipol==1: pl.xlabel(d.corrname[iprod])
      pl.suptitle(d.fname+' ringing index '+d.corrs[iprod]+' _'+title)
      pl.savefig(d.pldir+d.fname+'_ringind_'+d.corrs[iprod]+'_'+title+'.png', width=Inches(5.0))
  
      if d.doc:
        d.docfull.add_picture(d.pldir+d.fname+'_ringind_'+d.corrs[iprod]+'_'+title+'.png', width=Inches(5.0))


################################################################################################
def spectralperiods(d):
  """
  Take fft of each amp spectrum in scan averaged data and look for periodicities using fft after detrending. Note
  that max period detectable depends on detrending scale. 
  """
  
  # Look for periodicities in spectral domain using fft in auto and cross in scanav data
  t1 = time.time()
  fft = N.fft
  thresh = d.cfg.fft_thresh; win = d.cfg.fft_win
  ncorr = [d.nant, d.nbl]
  print("Searching for periodicities in individual spectra with threshold", thresh)
  if d.doc: 
      mylibs.bothdoc(d, "title", "Spectral periodicities")
      mylibs.bothdoc(d, "para", "Using threshold of "+str(thresh)+" sigma\n"+ "after detrending by polyfit("+str(win)+")")
  waves, phss = [], []
  bad = [[],[]]; badtitle = [[],[]]; badinds = [[],[]]
  for iprod, prod in enumerate(d.corrs):
      print("Processing", prod, "scan", end=' ')
      nn = N.zeros((d.nband,d.npol,len(d.gscans),ncorr[iprod]))
      wave, phs = N.zeros(nn.shape)*N.nan, N.zeros(nn.shape)*N.nan
      ctr = N.zeros(d.nband,int)
      for ii,scan in enumerate(d.gscans):
          mylibs.mywrite(scan)
          data = mylibs.myload(d.sdpdir+'scan_'+str(scan)+'/scan_'+str(scan)+'_'+prod+'_scanav')
          for iband in range(d.nband):
              bchan, echan = d.bands[iband]
              dum = N.abs(data[0][:,bchan:echan,:])
              for iflags in range(len(d.bandflags[iband])):
                  dum[:,d.bandflags[iband][iflags][0]-bchan:d.bandflags[iband][iflags][1]-bchan] = N.nan
              for ipol in range(d.npol):
                  for icorr in range(ncorr[iprod]):
                    y = dum[ipol,:,icorr]
                    y = y - mylibs.poly_filter(y, win)
                    nu, amp, angle, fy = mylibs.fft1(y)
  
                    med = N.nanmedian(amp)
                    mad = 1.5*N.nanmedian(N.abs(amp-med))
                    amp[:int(len(amp)/2)+int(d.cfg.fft_flag)] = 0; amp[-10:] = N.nan
                    inds = N.where((amp-med)/mad >thresh)[0]
  
                    if len(N.where(inds>len(amp)/2)[0]) > 0: 
                        nn[iband,ipol,ii,icorr] += 1
  
                    if len(inds)>0:
                        wave[iband,ipol,ii,icorr] = N.abs(1./nu[N.nanargmax(amp)])
                        phs[iband,ipol,ii,icorr] = N.abs(angle[N.nanargmax(amp)])
                        bad[iprod].append([y, nu, amp, fft.fftshift(N.abs(fy)), (N.nanmax(amp)-med)/mad])
                        badinds[iprod].append([iprod,scan,iband,ipol,icorr])
                        str1="Band"+str(iband)+" Scan"+str(scan)+" Pol"+str(ipol)+d.corrname[iprod]+str(icorr)
                        badtitle[iprod].append(str1)
                        ctr[iband] += 1
      print()
      for iband in range(d.nband):
          print("  Found", ctr[iband], " bad spectra in ", prod, "in band", iband)
          if d.doc: mylibs.bothdoc(d, "para", "Found "+str(ctr[iband])+" bad spectra in "+ prod+" in band"+str(iband),style="List Bullet")
  
      waves.append(wave); phss.append(phs)
      print("  Found", len(badtitle[iprod]), "bad spectra in ", prod)
  
  
  # Plot bad stuff
  #
  for iprod in range(len(d.corrs)):
      if len(bad[iprod])==0:
          print("No significant fft signal in any", d.corrs[iprod])
      else:
          print("Significant FFT peak found in", d.corrs[iprod], len(bad[iprod]))
          
  # Plot data and fft for auto
  nums = [min(int(d.cfg.fft_nmax),len(bad[0])), min(int(d.cfg.fft_nmax),len(bad[1]))]
  for jj in range(2):
      if nums[jj]>0:
        pl.figure(figsize=(8,12))
        for i in range(nums[jj]):
            pl.subplot(nums[jj],2,2*i+1)
            pl.plot(bad[jj][i][0])
            pl.title(badtitle[jj][i])
            pl.subplot(nums[jj],2,2*i+2)
            pl.plot(bad[jj][i][1],bad[jj][i][3])
            pl.plot(bad[jj][i][1],bad[jj][i][2])
            if i==0: pl.title('FFT')
      pl.suptitle(d.fname+' FFT for '+d.corrs[jj])
      pl.savefig(d.pldir+d.fname+'_spectralfft_'+d.corrs[jj]+'_examples.png')
      if d.doc:
          d.docfull.add_picture(d.pldir+d.fname+'_spectralfft_'+d.corrs[jj]+ '_examples.png', width=Inches(5.0))
  
  for iprod in range(len(d.corrs)):
      for iband in range(d.nband):
          nn = N.sum(~N.isnan(waves[iprod][iband]))
          print("For", d.corrs[iprod], "band", iband, "get", nn, "points")
  
  # Plot the amps of scan-chan for auto and cross
  for iprod in range(len(d.corrs)):
      for iband in range(d.nband):
          nn = N.sum(~N.isnan(waves[iprod][iband]))        
          if nn>0:
              pl.figure(figsize=(10,6)) 
              for ipol in range(d.npol):
                  pl.subplot(2,1,ipol+1); libs.imshow(N.transpose(waves[iprod][iband,ipol])); pl.colorbar(); 
                  pl.title('Band'+str(iband)+' Pol'+str(ipol)); pl.ylabel('Scan'); 
                  if ipol==1: pl.xlabel(d.corrname[iprod])
              pl.suptitle(d.fname+' Spectral FFT periods for '+d.corrs[iprod]+' band'+str(iband))
              pl.savefig(d.pldir+d.fname+'_spectralfft_periods_scanav_'+d.corrs[iprod]+'_band'+str(iband)+'.png')
              if d.doc:
                  d.docfull.add_picture(d.pldir+d.fname+'_spectralfft_periods_scanav_'+d.corrs[iprod]+'_band'+str(iband)+'.png', width=Inches(5.0))
  
  # Plot histo of periodicities for both
  for iprod in range(len(d.corrs)):
      if len(bad[iprod])>10:
          nus, amps = N.asarray(bad[iprod])[:,1], N.asarray(bad[iprod])[:,2]
          maxv = N.asarray(bad[iprod])[:,4]
          pers = 1./N.asarray([nus[i][N.nanargmax(amps[i])] for i in range(nus.shape[0])])
          pl.figure(figsize=(9,4))
          for iband in range(d.nband):
              inds = N.where(N.asarray(badinds[iprod])[:,2]==iband)
              pl.subplot(1,3,iband+1)
              dum = pl.hist(pers[inds],40)
              v1,v2 = N.unique(pers,return_counts=True)
              ind=N.argmax(v2)            
              pl.xlabel('Spectral FFT period')
              pl.title('Band'+str(iband))
              print("%s %s %i %s %.2f" %(d.corrs[iprod], ' band', iband, 'max at', v1[ind]))
          pl.suptitle(d.fname+' spec fft period for '+d.corrs[iprod])
          pl.savefig(d.pldir+d.fname+'_spectralfft_hist_period_scanav_'+d.corrs[iprod]+'.png')
          if d.doc:
              d.docfull.add_picture(d.pldir+d.fname+'_spectralfft_hist_period_scanav_'                        +d.corrs[iprod]+'.png', width=Inches(5.0))
  
  # Plot fft value versus period
  cols = ['b','r','g','m','k','c']
  for iprod in range(len(d.corrs)):
      if len(bad[iprod])>5:
          pl.figure()
          pers, maxv = [[],[],[]], [[],[],[]]
          for iband in range(d.nband):
              for ibad in range(N.asarray(bad[iprod]).shape[0]):
                  bad1 = N.asarray(bad[iprod])[ibad]
                  iiprod, iiscan, iiband, iipol, iicorr = N.asarray(badinds[iprod])[ibad]
                  if iiband==iband and iiprod==iprod:
                      nu, amp = N.asarray(bad[iprod])[ibad,1], N.asarray(bad[iprod])[ibad,2]
                      pers[iband].append(1./nu[N.nanargmax(amp)])
                      maxv[iband].append(N.asarray(bad[iprod])[ibad,4])
              pl.plot(pers[iband], maxv[iband], '.', color=cols[iband], label=d.corrs[iprod]+str(iband))
          pl.legend()
          pl.xlabel('Spectral FFT period'); pl.ylabel('FFT amplitude SNR')
          pl.title(d.corrs[iprod])
          pl.legend()
          pl.savefig(d.pldir+d.fname+'_spectralfft_period_fftamp_scanav_'+d.corrs[iprod]+'.png')
          if d.doc:
              d.docfull.add_picture(d.pldir+d.fname+'_spectralfft_period_fftamp_scanav_'+d.corrs[iprod]+'.png',
                      width=Inches(5.0))
  
  t2 = time.time()
  print("%s %.1f %s\n" %("Done in", t2-t1, "sec"))


################################################################################################

def get_fullband(d, x):

    bchan, echan = d.fullband
    for flag in d.fullbandflags:
        x[flag[0]:flag[1]] = N.nan

    return x[bchan:echan]

################################################################################################

def detrend(d):
  """
  Detrend each spectrum in scan averaged data for each scan by polynomial filtering over window 
  (default 31). Poly filter is much better than median filter.
  """
  t1 = time.time()
  print("Detrend each scan averaged spectrum")
  fft = N.fft
  win = d.cfg.detrend_win
  bchan, echan = d.fullband
  d.docfull.add_heading("Detrending")
  d.docfull.add_paragraph("Applying polynomial filter with window "+str(win)+" channels")
  ncorr = [d.nant, d.nbl]
  d.pf_all = []; d.pf_full = []
  for iprod, prod in enumerate(d.corrs):
      d.pf_data = N.zeros((d.npol, d.gnscan, echan-bchan, ncorr[iprod]))
      d.pf_1 = N.zeros((d.npol, d.gnscan, d.nchan, ncorr[iprod]))
      print("Processing", prod, "scan", end=' ')
      for ii,scan in enumerate(d.gscans):
          mylibs.mywrite(scan)
          data = mylibs.myload(d.sdpdir+'scan_'+str(scan)+'/scan_'+str(scan)+'_'+prod+'_scanav')
          dum = N.abs(data[0]) # npol,nchan,ncorr
          for ipol in range(d.npol):
              for icorr in range(ncorr[iprod]):
                  x = get_fullband(d, dum[ipol,:,icorr])
                  norm = 1.0
                  if iprod==0: norm = N.nanmean(x)
                  x = x - mylibs.poly_filter(x, win)
                  d.pf_data[ipol,ii,:,icorr] = x/norm
  
                  x = N.abs(dum[ipol,:,icorr])
                  d.pf_1[ipol,ii,:,icorr] = (x-mylibs.poly_filter(x, win))/norm
      print()
      d.pf_all.append(d.pf_data)
      d.pf_full.append(d.pf_1)

  t2 = time.time()
  print("%s %.1f %s\n" %("Done in", t2-t1, "sec"))

################################################################################################

def ringing_2chan_single(d):

  """
  See if there is 2-channel spectral ringing in the ant/bl-scan averaged spectra for the full band
  """
  t1 = time.time()
  print("Look for 2-channel ringing in the global average detrended spectrum")
  if d.doc: mylibs.bothdoc(d, "title", "2-chan ringing in average detrended spectrum")
  print("2-channel ringing for corr-scan averaged pf spectrum; ringing index for")
  for iprod, prod in enumerate(d.corrs):
    for ipol in range(d.npol):
      avspec = N.nanmean(N.nanmean(d.pf_all[iprod][ipol],0),1)
      rind = mylibs.chan2ripple(avspec, 0, domf=False, win=d.cfg.ring_win)
      str1 = "%s %s %s %i %s %.2f" %("Ringing index for", prod, " pol", ipol, "is ", rind)
      print(str1)
      if d.doc:
          mylibs.bothdoc(d, "para", str1, style="List Bullet")
  t2 = time.time()
  print("%s %.1f %s\n" %("Done in", t2-t1, "sec"))

################################################################################################

def xycorrcoeff_spec(d):
  """
  Calculate corr coeff between X and Y for each detrended scanav spectrum
  """

  from scipy.stats import pearsonr
  t1 = time.time()
  print("Calculate corr coeff bvetween X and Y pol for each detrended scanav spectrum")
  # Corr coeff of average spectra
  if d.doc: mylibs.bothdoc(d, "title", "X-Y pol correlation in spectra")
  pl.figure(figsize=(9,4))
  corrcoefs = []; ncorr = [d.nant, d.nbl]
  for iprod, prod in enumerate(d.corrs):
      corrcoef = N.zeros((d.gnscan,ncorr[iprod]))
      for iscan in range(d.gnscan):
          for icorr in range(ncorr[iprod]):
              xx = d.pf_all[iprod][0,iscan,:,icorr]
              yy = d.pf_all[iprod][1,iscan,:,icorr]
              xx = mylibs.madflag(xx, thresh=d.cfg.cc_thresh)
              yy = mylibs.madflag(yy, thresh=d.cfg.cc_thresh)
              inds = N.where(~N.isnan(xx*yy))[0]
              xx = xx[inds]; yy = yy[inds]
              if len(xx)>d.cfg.cc_minlen:
                corrcoef[iscan,icorr] = pearsonr(xx,yy)[0]
              else:
                corrcoef[iscan,icorr] = N.nan
      corrcoefs.append(corrcoef)
  
      for iscan in range(d.gnscan):
          pl.subplot(1,2,iprod+1)
          pl.plot(corrcoef[iscan],'.')
          pl.title(prod+' corrcoeff of pol of spectra')
          pl.xlabel(d.corrname[iprod])
  pl.suptitle(d.fname+' corr coeff for X vs Y for each corr-scan spectrum')
  pl.savefig(d.pldir+d.fname+'_corrcoeff_pol_indspec.png')
  if d.doc: d.docfull.add_picture(d.pldir+d.fname+'_corrcoeff_pol_indspec.png', width=Inches(5.0))
  
  str1 = "%s %.2f %.2f" %("  Mean and median corrcoeff for auto is  ", N.nanmean(corrcoefs[0]), N.nanmedian(corrcoefs[0]))
  str2 = "%s %.2f %.2f" %("  Mean and median corrcoeff for cross is ", N.nanmean(corrcoefs[1]), N.nanmedian(corrcoefs[1]))
  print(str1); print(str2)
  if d.doc: mylibs.bothdoc(d, "para", [str1, str2])
  
  pl.figure(figsize=(9,7))
  for iprod,prod in enumerate(d.corrs):
      pl.subplot(2,1,iprod+1); libs.imshow(N.transpose(corrcoefs[iprod]))
      pl.ylabel('Scan'); pl.xlabel(d.corrname[iprod]); pl.title(prod); pl.colorbar()
  pl.suptitle(d.fname+' corr coeff for X vs Y for each corr-scan spectrum')
  pl.savefig(d.pldir+d.fname+'_corrcoeff_pol_image.png')
  if d.doc: d.docfull.add_picture(d.pldir+d.fname+'_corrcoeff_pol_image.png', width=Inches(5.0))
  
  pl.figure(figsize=(9,8))
  fold = 64
  for iprod, prod in enumerate(d.corrs):
      av0 = N.nanmean(N.nanmean(d.pf_all[iprod][0],0),1)
      av1 = N.nanmean(N.nanmean(d.pf_all[iprod][1],0),1)
      av0 = mylibs.madflag(av0, thresh=d.cfg.cc_thresh)
      av1 = mylibs.madflag(av1, thresh=d.cfg.cc_thresh)
      fold0 = N.asarray([N.nanmean(av0[i::fold]) for i in range(fold)])
      fold1 = N.asarray([N.nanmean(av1[i::fold]) for i in range(fold)])
      pl.subplot(2,2,iprod*2+1);pl.plot(av0,av1,'.'); pl.title(prod+'av spec')
      pl.subplot(2,2,iprod*2+2);pl.plot(fold0,fold1,'.'); pl.title(prod+'fold spec')
      inds = N.where((~N.isnan(av0))*(~N.isnan(av1)))
      str1 = "%s %s %s %.2f" %("  Corr coeff for",prod,"av spectrum X vs Y is   ", pearsonr(av0[inds],av1[inds])[0])
      inds = N.where((~N.isnan(fold0))*(~N.isnan(fold1)))
      str2 = "%s %s %s %.2f" %("  Corr coeff for",prod,"folded av spec X vs Y is", pearsonr(fold0[inds],fold1[inds])[0])
      print(str1); print(str2)
      if d.doc: mylibs.bothdoc(d, "para", str1, style="List Bullet")
      if d.doc: mylibs.bothdoc(d, "para", str2, style="List Bullet")
  pl.suptitle(d.fname+' corr coeff for X vs Y for average spectrum')
  pl.savefig(d.pldir+d.fname+'_corrcoeff_pol_avspec.png')
  if d.doc: d.docfull.add_picture(d.pldir+d.fname+'_corrcoeff_pol_avspec.png',width=Inches(5.0))

  t2 = time.time()
  print("%s %.1f %s\n" %("Done in", t2-t1, "sec"))


################################################################################################

def check_64fold(d):
  """
  Fold every 64 chans (and multiples) and plot spectra
  """
  t1 = time.time()
  print("Fold the spectra to locate 64-channel dips")
  if d.doc: d.docfull.add_heading("64-channel folded spectra")
  for ii in range(2):
    if d.gnscan<7: dum4=d.gnscan
    else: dum4=1
    for isc in range(dum4):
      pl.figure(figsize=(14,10))
      for ipol in range(d.npol):
          if dum4==1:
              av = N.nanmean(N.nanmean(d.pf_all[ii][ipol],0),1)
          else:
              av = N.nanmean(d.pf_all[ii][ipol][isc],1)
          for ifold,fold in enumerate([64, 128, 256]):
              folded_sp = N.asarray([N.nanmean(av[i::fold]) for i in range(fold)])
              pl.subplot(3,2,2*ifold+1); pl.plot(folded_sp)
              pl.title(' fold'+str(fold))

      pl.subplot(1,2,2)
      fold = 128
      cols = ['steelblue','darkorange','c','k']
      for ipol in range(d.npol):
          if dum4==1:
              av = N.nanmean(N.nanmean(d.pf_all[ii][ipol],0),1)
          else:
              av = N.nanmean(d.pf_all[ii][ipol][isc],1)
          folded_sp = N.asarray([N.nanmean(av[i::fold]) for i in range(fold)])
          if ipol==0: fact = N.nanmax(N.abs(folded_sp))*1.2
          for i in range(int(fold/64)):
              pl.plot(folded_sp[i*64:(i+1)*64]+i*fact, color=cols[ipol])
              pl.plot(folded_sp[i*64:(i+1)*64]+i*fact, color=cols[ipol])
      pl.title('Fold-128 in two parts offset')
      pl.suptitle(d.fname+' '+d.corrs[ii]+' fold 64n of spec av-scan-corr band4')
      pl.savefig(d.pldir+d.fname+'_'+d.corrs[ii]+'_fold64n_avspec.png')
      if d.doc: d.docfull.add_picture(d.pldir+d.fname+'_'+d.corrs[ii]+'_fold64n_avspec.png',width=Inches(5.0))
  
  pl.figure(figsize=(10,int(d.gnscan*1.5)))
  for isc in range(d.gnscan):
      for ipol in range(d.npol):
          av = N.nanmean(d.pf_all[1][ipol][isc],1)
          fold = 128
          folded_sp = N.asarray([N.nanmean(av[i::fold]) for i in range(fold)])
          pl.subplot(d.gnscan,2,isc*2+1); pl.plot(folded_sp)
          pl.subplot(d.gnscan,2,isc*2+2); pl.plot(av)
  pl.savefig(d.pldir+d.fname+'_'+d.corrs[1]+'_fold_perscan_avspec.png')
  
  pl.figure()
  for ipol in range(d.npol):
      av = N.nanmean(N.nanmean(d.pf_all[1][ipol],0),1)
      fsp = N.asarray([N.nanmean(av[i::128]) for i in range(128)])
      pl.subplot(211); pl.plot(av)
      pl.subplot(212); pl.plot(fsp)
  pl.savefig(d.pldir+d.fname+'_'+d.corrs[1]+'_ffold_perscan_avspec.png')
  if d.doc: d.docfull.add_picture(d.pldir+d.fname+'_'+d.corrs[1]+'_ffold_perscan_avspec.png',width=Inches(5.0))

  t2 = time.time()
  print("%s %.1f %s\n" %("Done in", t2-t1, "sec"))

################################################################################################

def do_comb(d):
  """
  Do comb analysis for all start,fold combinations
  """

  t1 = time.time()
  print("Comb analysis for corr-scan averaged pf spectra")
  if d.doc:
      mylibs.bothdoc(d, "title", "Comb function analysis")
      d.docfull.add_paragraph("Using nsearch="+str(d.cfg.comb_nsearch)+"; minfold="+str(d.cfg.comb_minfold)+\
                              "; thresh="+str(d.cfg.comb_thresh))
  
  for iprod, prod in enumerate(d.corrs):
      for ipol in range(d.npol):
          print(prod, "and pol", ipol)
          avspec = N.nanmean(N.nanmean(d.pf_all[iprod][ipol],0),1)
          arr, mx = mylibs.comb_analysis(avspec, nmin=d.cfg.comb_mnn, nmax=d.cfg.comb_nmax)
          str1, str2 = mylibs.clean_comb(arr, d.cfg.comb_mnn, d.cfg.comb_nsearch, d.cfg.comb_minfold, d.cfg.comb_thresh, \
                       doplot=True, title=d.fname+' '+prod+' pol'+str(ipol)+' avspec',\
                       savefig=d.pldir+d.fname+'_comb_'+prod+'_pol'+str(ipol)+'_avspec.png')

          if d.doc:
              mylibs.bothdoc(d, "para", str1, style="List Bullet")
              d.docfull.add_picture(str2,width=Inches(5.0))

  t2 = time.time()
  print("%s %.1f %s\n" %("Done in", t2-t1, "sec"))

################################################################################################

def thomas_64(d):
  """
  Do the 64 chan dip search using Thomas's method
  """

  t1 = time.time()
  print("Search for 64 channel dip using Thomas's method")
  corr = ['Auto', 'Cross']
  pl.figure(figsize=(10,8))
  for iprod in range(2):
      for ipol in range(d.npol):
          avspec = N.nanmean(N.nanmean(d.pf_full[iprod][ipol],0),1)
  
          renorm=1000.0; minstep=8
          pl.subplot(2,2,iprod*2+ipol+1)
          for a in N.arange(0,4096,minstep):
              b=1
              idx=abs(N.arange(d.nchan)-a).argmin()
              while (a%(minstep*2**b)==0)&(b<13): b=b+1
              try:
                if abs(avspec[idx-1]-avspec[idx+1])<0.005:
                  dipdepth = -(avspec[idx-1]+avspec[idx+1])*0.5+avspec[idx]
                  pl.plot((b-1)+N.random.randn()*0.07,dipdepth*renorm,ms=14,\
                          alpha=N.clip((b/10.)**1.6,0,1),marker='.',color='m')
              except:pass
          pl.xticks(N.arange(12),minstep*2**N.arange(12));
          pl.grid(); pl.ylim([-4,1]); pl.xlim([-0.2,10.2]);
          pl.title(corr[iprod]+' pol'+str(ipol))
  
  pl.savefig(d.pldir+d.fname+"_2powerplot.png")
  if d.doc:  d.docfull.add_picture(d.pldir+d.fname+"_2powerplot.png",width=Inches(5.0))
  t2 = time.time()
  print("%s %.1f %s\n" %("Done in", t2-t1, "sec"))

################################################################################################


