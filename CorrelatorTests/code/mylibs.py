#pairs) My functions for analysing data which are numpy arrays

#from firstcell import *
import numpy as N
from scipy.ndimage.filters import median_filter
from math import *
import os, sys, glob, libs, warnings, time
warnings.filterwarnings("ignore")

################################################################################################
def factors(n):
    """ All factors of a number in desc order except 1 and itself"""
    dum = []
    for i in range(n-1, 1, -1):
        if n%i==0: dum.append(i)
    dum = N.asarray(dum)

    return dum

################################################################################################

def mywrite(num,dum=None):
    """ Pretty write for ongoing process"""
    print(num, end=', '); sys.stdout.flush()

################################################################################################

def madflag(arr, thresh=5.0):
    """ Flag based on MAD and return flagged array"""

    med = N.nanmedian(arr)
    mad = 1.5*N.nanmedian(N.abs(arr-med))
    arr1 = N.where(N.abs(arr-med)/mad>thresh, N.nan, arr)

    return arr1

################################################################################################

def ddir(dname):

    return dname+'/'

################################################################################################

def getmad(x, axis=1, robust=True, niter=1, doabs=True, thresh=5.0):
    """
    Calculate 1.5*MAD. If x is more than 1d, do it on axis.
    Do this for niter iterations after throwing out points > thresh
    doabs means take N.abs(y-med)/mad 
    Robust is false => use mean,std else use med,mad
    """

    if doabs: func = N.abs
    else: func = N.asarray

    y = N.copy(x)*1.0
    med, mad = None, None
    if x.ndim==1:
        for i in range(niter):
            if robust: 
                med = N.nanmedian(y)
                mad = 1.5*N.nanmedian(N.abs(y-med))
            else:
                med = N.nanmean(y)
                mad = N.nanstd(y)

            ind = N.where(func(y-med)/mad>thresh)[0]
            y[ind] = N.nan

    if x.ndim==2:
        for i in range(niter):
            if robust:
                med = N.nanmedian(y, axis)
                mad = 1.5*N.nanmedian(N.abs(y-N.stack([med]*y.shape[axis],axis)), axis)
            else:
                med = N.nanmean(y, axis)
                mad = N.nanstd(y, axis)
            ind = N.where(func(y-N.stack([med]*y.shape[axis],axis))/N.stack([mad]*y.shape[axis],axis)>thresh)
            y[ind] = N.nan

    return med, mad, (N.nanmin(y), N.nanmax(y))


################################################################################################

def poly_filter(x, win, order=5, beg=0, end=None):
    """
    Median filter is bad when data is monotonic even with noise, since the filtered values
    are then exactly equal to the data. ploy fit seems better. See ng_methods.png
    """

    n = len(x)
    if end==None: end = n
    op = N.zeros(n)
    for j in range(beg,end,win):
        n1 = j+win
        if j+win>n: n1 = end
        x1 = N.arange(j,n1); y1 = x[j:n1]
        inds = N.where(~N.isnan(y1))[0]
        if len(inds)>order:
          z = N.polyfit(x1[inds], y1[inds], order)
          dum = N.poly1d(z)
          dum1 = dum(N.arange(j,n1))
          op[j:n1] = dum1
        else:
          op[j:n1] = N.nan

    return op
                
################################################################################################

def fft1(y):
    fft = N.fft
    inds = N.isnan(y)
    y[inds] = 0.

    fy = fft.fft(y)
    ff = fft.fftfreq(len(y))
    nu = fft.fftshift(ff)
    amp = fft.fftshift(N.abs(fy))
    real = fft.fftshift(N.real(fy))
    imag = fft.fftshift(N.imag(fy))
    angle = fft.fftshift(N.angle(fy, 1))

    return nu, amp, angle, fy

################################################################################################

def myload(fname):
    import pickle

    if os.path.isfile(fname+'.pickle'):
        return pickle.load(open(fname+'.pickle', "rb"), encoding='bytes')
    else:
        if os.path.isfile(fname+'.npy'):
            return N.load(fname+'.npy')
        else:
            raise RuntimeError("No file found as "+fname)


################################################################################################
def create_obj(num):

    class myclass(object):
        pass

    basedir = '/data/mohan/'
    basedir += '/'
    
    d = myclass()

    if isinstance(num,int): num = str(num)
    fname = num
    create_dirs(d, fname, basedir)

    return d

################################################################################################

def create_dirs(d, fname=None, basedir=None):
    """
    For a fname (id of file), create _pickle dir and dirs
    under it, if they dont exist
    """
    import glob

    if basedir==None: 
        if hasattr(d, 'basedir'): basedir = d.basedir
        else: basedir = ddir('/data/mohan/')
    if fname==None: 
        if hasattr(d, 'fname'): fname = d.fname
        else: raise RuntimeError("Need fname")
    else:
        d.fname = fname

    fdir = ddir(basedir+fname)
    if not os.path.isdir(fdir):
        os.mkdir(fdir)
        print("Made directory", fdir)
    d.fdir = ddir(basedir+fname)

    sdpdir = ddir(fdir+fname+'_pickle/')
    if not os.path.isdir(sdpdir):
        os.mkdir(sdpdir)
        print("Made directory", sdpdir)
    d.sdpdir = ddir(sdpdir)

    casadir = ddir(fdir+'/')
    #if not os.path.isdir(casadir):
    #    os.mkdir(casadir)
    #    print("Made directory", casadir)
    d.casadir = ddir(casadir)

    pldir = ddir(fdir+'/plots/')
    if not os.path.isdir(pldir):
        os.mkdir(pldir)
        print("Made directory", pldir)
    d.pldir = ddir(pldir)

    caldir = ddir(fdir+'/calib/')
    if not os.path.isdir(caldir):
        os.mkdir(caldir)
        print("Made directory", caldir)
    d.caldir = ddir(caldir)

################################################################################################

def dodocs(d):
    from docx import Document

    d.docsumm = Document()
    d.docsummname = d.fdir+d.fname+'_summaryreport.docx'
    d.docsumm.add_heading(d.fname+' - summary report', 0)

    d.docfull = Document()
    d.docfullname = d.fdir+d.fname+'_fullreport.docx'
    d.docfull.add_heading(d.fname+' - full report', 0)

################################################################################################

def bothdoc(d, mode, str1, style=None):

    if d.doc:
        if mode=='title':
            d.docfull.add_heading(str1)
            d.docsumm.add_heading(str1)
        if mode=='para':
            if isinstance(str1, str): str1 = [str1]
            for str2 in str1:
              d.docfull.add_paragraph(str2, style=style)
              d.docsumm.add_paragraph(str2, style=style)

################################################################################################
def read_katdal_para(f, d, flog=None):

    f.select()
    d.nant = d.na = d.nants = len(f.ants)
    d.nt, d.nchan, d.ncorr = f.shape
    d.nbl = int(d.na*(d.na-1)/2)
    d.nscan = len(f.scan_indices)
    d.nch = int(d.nchan/1024)
    band = f.spectral_windows[0].band
    if band=='L': d.band='l'
    else: d.band='u'

    print('Shape ', f.shape)
    print("%-10s %i" %('Num ants ', d.na))
    print("%-10s %i" %('Num chans ', d.nchan))
    print("%-10s %i" %('Num corr', d.ncorr))
    print("%-10s %i" %("Num scans ", d.nscan))
    if int((d.nant*(d.nant-1)/2+d.nant)*4) != d.ncorr: raise RuntimeError("Wrong indexes")

    if not hasattr(d, 'doc'): d.doc = False
    if d.doc:
        for dd in [d.docfull, d.docsumm]:
            dd.add_heading('Data parameters', level=1)
            dd.add_paragraph('Shape : '+str(f.shape), style='List Bullet')
            dd.add_paragraph("%-15s %i" %('Num of ants ', d.na), style='List Bullet')
            dd.add_paragraph("%-15s %i" %('Num of corr', d.ncorr), style='List Bullet')
            dd.add_paragraph("%-15s %i" %("Num of chans ", d.nchan), style='List Bullet')
            dd.add_paragraph("%-15s %i" %("Num of scans ", d.nscan), style='List Bullet')

################################################################################################

def read_bptable(fname, bpname):
    """ When you run this from the directory itself """
    import glob
    import katdal

    basedir = '/data/mohan/'
    class myclass(object):
           pass
    d = myclass()

    d.fname = fname
    f = katdal.open(basedir+'/'+fname+'/'+d.fname+'_sdp_l0.rdb')
    read_katdal_para(f, d)
    
    read_casa_bp(d, bpname)
    print(d.data.shape)
    reshape_bptable(d)
    print(d.data.shape)
    
    return d.data
    
################################################################################################

def read_sntable(tname):
    tb.open(tname)
    d = N.squeeze(tb.getcol('CPARAM'))
    time = tb.getcol('TIME')
    ntime = len(N.unique(time))
    na = d.shape[1]/ntime
    if na*ntime != d.shape[1]: raise RuntimeError("Cant decompose into matrix")
    d = d.reshape((d.shape[0], ntime, na))
    return d
    

################################################################################################

def read_casa_bp(d, bpname):
    """ Read bandpass table written out by CASA"""

    import casacore.tables as t
    mydir = ddir(d.casadir+bpname)
    if not os.path.isdir(mydir):
        raise RuntimeError(mydir+" does not exist")

    tt = t.table(mydir)

    d.times = tt.getcol('TIME')
    d.ant1 = tt.getcol('ANTENNA1')
    d.ant2 = tt.getcol('ANTENNA2')
    d.scannum = tt.getcol('SCAN_NUMBER')
    d.data = tt.getcol('CPARAM')
    d.err = tt.getcol('PARAMERR')
    d.flag = tt.getcol('FLAG')
    d.snr = tt.getcol('SNR')

    d.ncorr, d.nchan, d.npol = d.data.shape
    d.nscan = len(N.unique(d.scannum))
    d.nant = len(N.unique(d.ant1))
    d.ants = N.unique(d.ant1)
    
    if d.data.shape[0] == d.ncorr: print("*** Need to reshape data array")

################################################################################################
def reshape_bptable(d, do=False):
    """ 
    Reshape complex data & flag from (corr X chan X pol) to (pol X scan X ant X chan)
    and separate out amp and ph to give bpamp bpph, each is (pol X scan X ant X chan)
    """

    def polfirst(arr):

        dum = []
        for i in range(arr.shape[-1]): # npol
            dum.append(arr[...,i])

        return N.asarray(dum)

    nscan, nant, nchan, npol = d.nscan, d.nant, d.nchan, d.npol
    print('nscan, nant, nchan, npol', d.nscan, d.nant, d.nchan, d.npol)

    if d.data.shape[0] == d.ncorr:
      d.data = d.data.reshape((nscan, nant, nchan, npol))
      d.err = d.err.reshape((nscan, nant, nchan, npol))
      d.snr = d.snr.reshape((nscan, nant, nchan, npol))
      d.flag = d.flag.reshape((nscan, nant, nchan, npol))
      d.flag = N.where(d.flag,N.nan,1)

      d.data = polfirst(d.data)
      d.err = polfirst(d.err)
      d.snr = polfirst(d.snr)
      d.flag = polfirst(d.flag)

    if (d.data.shape[0] == d.ncorr) or do:
      # Write to better vars, now npol X nscan X nant X nchan
      d.bpamp = N.abs(d.data)*d.flag
      d.bpph = N.angle(d.data)*d.flag

    else:
      print("*** No need to reshape or pol-unpack data array")


################################################################################################
def timeavg_comp(vis, flags, na, nvis, npol, nchan):
    """
    Returns scan averaged data to be pickled later.

    For cross, returns npol*[mean, std, dc] where mean and std (complex)
    are over the scan, after applying sdp flags
    The dc is averaged over the scan without appling flags
    For auto, returns npol*[mean, std, dc] and is same as cross
    Autocorr for a pol is [ntime, nchan, nant] and cross is [ntime, nchan, nbl]

    Doing average on complex values directly since I assume reasonable phase stability within a scan
    ** Do median by hand since there is some issue with median of complex arrays if needed
    ** med = N.nanmedian(N.real(cross),0)+N.nanmedian(N.imag(cross))*1j
    """
    import pylab as pl
    pl.ion()

    retc, reta = [], []
    nbl = int(na*(na-1)/2)
    # Do cross
    mean = N.zeros((npol,nchan,nbl))*1j
    std = N.zeros((npol,nchan,nbl))*1j
    dc = N.zeros((npol,nbl))*1j
    for i in range(npol):
        cross = vis[:,:,2*na+i*nvis:2*na+(i+1)*nvis]
        if flags is not None:
          flag = flags[:,:,2*na+i*nvis:2*na+(i+1)*nvis]   # only reading hh,vv so 4->2

        dc[i] = N.nanmean(cross[:,0,:],0) # av over scan, is just nant for a pol
        if flags is not None:
          cross = cross*flag
          del flag

        mean[i] = N.nanmean(cross,0)
        std[i] = N.nanstd(cross,0)
        del cross

    retc = mean, std#, dc
    del mean, std
    print("  Done with cross")
        
    # Do auto
    mean = N.zeros((npol,nchan,na))*1j
    std = N.zeros((npol,nchan,na))*1j
    dc = N.zeros((npol,na))*1j
    for i in range(npol): # npol pols
        auto = N.abs(vis[:,:,i*na:(i+1)*na])
        if flags is not None:
          flag = flags[:,:,i*na:(i+1)*na]

        dc[i] = N.nanmean(auto[:,0,:],0)  # av over scan, is just nbl for a pol
        if flags is not None:
          auto = auto*flag

        mean[i] = N.nanmean(auto,0)
        std[i] = N.nanstd(auto,0)

    reta = mean, std# , dc   # so that N.save is faster

    return retc, reta

################################################################################################

def get_params(fname, pdir, suffix='_sdp_l0.full.rdb', fdir=None, output=True):
    """
    Open katdal and get basic params. Store in file called fname.listobs
    """

    if fdir==None: fdir ='/data/mohan/'
    fdir += '/'

    t1 = time.time()
    d = katdal.open(fdir+fname+suffix)
    t2 = time.time()
    print("Opened file in", (t2-t1)/60., "min")

    nant = na = nants = len(d.ants)
    nt, nchan, ncorr = d.shape
    nvis = na*(na-1)/2
    nscans = len(d.scan_indices)

    d.select(corrprods='cross', pol='hh,vv', scans='track')
    nscan = len(d.scan_indices)
    d.select()

    if output:
        print('  Shape ', (nt, nchan, ncorr))
        print('  Num ants  ', na)
        print('  Num chans ', nchan)
        print('  Num corr  ', ncorr)
        print("  Tot scans ", nscans)
        print("  Trk scans ", nscan)

    fn1 = pdir+fname+'.listobs'
    fn = open(fn1, "w")
    fn.write("%s %i\n" %("nant ", nant))
    fn.write("%s %i\n" %("nchan ", nchan))
    fn.write("%s %i\n" %("ncorr ", ncorr))
    fn.write("%s %i\n" %("nt ", nt))
    fn.write("%s %i\n" %("nscan ", nscan))
    fn.write("%s %i\n" %("nscans ", nscans))
    fn.close()

    print("Wrote ", fn1)

################################################################################################

def madfilt(x, win, thresh):
    """
    Nan all points above threshold using mad-med
    """
    
    x1 = median_filter(x, win, mode='nearest')
    myflags = N.ones(len(x))
    diff = x-x1
    med = N.nanmedian(diff)
    mad = 1.5*N.nanmedian(N.abs(diff-med))
    ind = N.where(N.abs(diff-med)/mad>thresh)[0]
    myflags[ind] = N.nan
    
    return myflags

################################################################################################

def flag_sp(x, win=None, thresh=None):
    """
    Flag based on multiple median filtering
    """
    if win==None or thresh==None:
      win = [1001, 301]
      thresh = [10., 5.]
    else:
      if len(win) != len(thresh):
          raise RuntimeError("Windows and thresholds not same length")
      else:
          nn = len(win)

    flag = N.ones(len(x))
    for i in range(nn):
      flag = flag * madfilt(x, win[i], thresh[i])

    return flag

################################################################################################

def outlier(arr, mode, thresh, doabs=True):
    """ 
    calc std/mad of arr and send back inds that are more than thresh
    mode=normal => mean, std and mode=robust is med, mad
    """

    if mode=='normal': 
        mean, std = N.nanmean(arr), N.nanstd(arr)

    if mode=='robust': 
        mean = N.nanmedian(arr)  # i know ...
        std = 1.5*N.nanmedian(N.abs(arr-mean))

    dum = (arr-mean)/std
    if doabs: dum = N.abs(dum)
    inds = N.where(dum>=thresh)

    return inds

################################################################################################

def comb_analysis(x, nmin=2, nmax=200, doerr=None, norm=True):
    """
    Do comb analysis from a 1d series by calculating values in fold, start space
    Fold goes from nmin to nmax
    If doerr is int, then average from i-doerr to i+doerr+1 else
    do just for each point
    """

    if doerr==None:
        doerr=0

    arr = N.zeros((nmax,nmax))
    # If x is all NaN then arr still has zero in its upper triang
    if N.sum(N.isnan(x))==len(x):
        return arr*N.nan, (N.nan, N.nan)
    
    for fold in range(nmin, nmax):
        for start in range(fold+1):
            if doerr==0:
                #xx = N.nanmean(N.abs(x[start::fold]))
                xx = N.abs(N.nanmean((x[start::fold])))
            else:        
                dum = []
                for jj in range(start, len(x), fold):
                    dum.append(N.nanmax(N.abs(x[max(0,jj-doerr):jj+doerr+1])))
                dum = N.asarray(dum)
                xx = N.nanmean(dum)
            n1 = len(x[start::fold])-N.sum(N.isnan(x[start::fold]))
            if not norm: n1 = 1.0
            try:
                arr[fold-nmin, start] = xx*sqrt(n1)
            except:
                print("CANT DO", xx*sqrt(n1), fold, nmin, start)

    return arr, N.unravel_index(N.nanargmax(arr), arr.shape)

################################################################################################

def comb_analysis_float(x, fold0, zoom, dfold=4):
    """
    Do comb analysis from a 1d series by calculating values in fold, start space
    Modified to search around a fold, for a given zoom
    """
    
    # If x is all NaN then arr still has zero in its upper triang
    if N.sum(N.isnan(x))==len(x):
        return N.nan, N.nan
    
    df = dfold*zoom 
    folds = N.arange(fold0-dfold, fold0+dfold, 1./zoom)
    arr = N.ones((len(folds),int(round(N.max(folds)+1))))*N.nan

    for ii, fold in enumerate(folds):
        for start in N.arange(int(fold)+1):
            xr = N.asarray(N.arange(start,len(x),fold),int)
            if xr[-1]>=len(x): xr = xr[:-2]  # floating pt issues since fold is float
            xx = N.abs(N.nanmean(x[xr]))

            n1 = len(xr)-N.sum(N.isnan(x[xr]))
            arr[ii,start] = xx*sqrt(n1)
    mx = N.unravel_index(N.nanargmax(arr), arr.shape)
    med = N.nanmedian(arr); mad = 1.5*N.nanmedian(N.abs(arr-med))
    maxv = N.nanmax(arr)

    return N.asarray([folds[mx[0]], mx[1]]), N.nanmax(arr), maxv/mad

################################################################################################


def comb_analysis_zoom(x, nmin=2, nmax=100,minsnr=15, zooms=[20,100]):
    """
    Do comb analysis from a 1d series by calculating values in fold, start space
    Fold goes from nmin to nmax
    If doerr is int, then average from i-doerr to i+doerr+1 else
    do just for each point
    Modify to search for non-integral fold periods. By zooming in.
    """

    if nmax==None: nmax=100
    # First get max for integer values
    arr, mx = comb_analysis(x, nmin, nmax)
    arr = N.where(arr==0,N.nan,arr)
    snr0 = N.nanmax(arr)/(1.5*N.nanmedian(N.abs(arr-N.nanmedian(arr))))

    fold0, start0 = mx
    fold0 = fold0+2

    if fold0 > 2 and snr0>minsnr:
        # number of harmonics to search for
        folds = [fold0]
        for ii in range(2,6):
            folds.append(int(fold0/ii))
            if folds[-1] < 8: break

        # Zoom in for each for each fac
        mxs, maxv = N.zeros((len(zooms),len(folds),2)), N.zeros((len(zooms),len(folds)))
        snr = N.zeros((len(zooms),len(folds)))
        for iz,zoom in enumerate(zooms):
            for ii,fold in enumerate(folds):
                mxs[iz,ii], maxv[iz,ii], snr[iz,ii] = comb_analysis_float(x, fold, zoom)

        mx = N.unravel_index(N.nanargmax(maxv), maxv.shape)

        return mxs[mx], (snr0,snr[mx])
    else:
        return N.asarray([N.nan, N.nan]), (N.nan, N.nan)


################################################################################################
def clean_comb(arr, mnn=2, nsearch=1, minfold=5, thresh=5.0, doplot=False, title='', savefig=''):
    """
    Search for nsearch consecutive maxima and deconvolve their response from 
    the fold-start array
    Goes down to minfold (while removing effects of all factors of original maxfold)
    """

    import pylab as pl
    pl.ion()
    maxs = []
    npol, size = arr.shape[:2]
    foldarr = N.copy(arr)
    dum = N.where(arr==0,N.nan,arr)

    for i in range(nsearch):  
      mx = N.unravel_index(N.nanargmax(foldarr), foldarr.shape)
      maxs1 = [mx[0]+mnn, mx[1]]
      maxs.append(maxs1)
      for k in range(0,size):
          for j in range(mx[0],size,mx[0]+mnn):
              try:
                  foldarr[j, mx[1]+(mx[0]+mnn)*k]=0
              except:
                  pass

      med = N.nanmedian(dum); mad = 1.5*N.nanmedian(N.abs(dum-med))
      maxsnr = (N.nanmax(dum)-med)/mad
      str1 = "%s %s %s %.2f" %("  Max (fold,start)", str(maxs1), " with snr ", maxsnr)
      print(str1)

      plots = 1
      if maxsnr >= thresh:
        plots = 2
        facts = factors(mx[0]+mnn)  # exclude 4 2
        facts = facts[N.where(facts>=minfold)[0]]
        #print("Used factors ", facts)
        for fact in facts:  # each factor of 128 down to >1 is new fold
            for k in range(-int(mx[1]/fact),int(size/fact)+1):
                for j in range(fact-mnn,size,fact):
                    try:
                        foldarr[j, mx[1]+fact*k]=0
                    except:
                        pass
                  
    if doplot:
      f1 = pl.figure(figsize=(5*plots,4))
      pl.subplot(1,plots,1)
      libs.imshow(dum, extent=[mnn-0.5,size+mnn-0.5,-0.5,size-0.5])
      pl.title('Input'); pl.xlabel('Fold')
      if plots==2:
        pl.subplot(1,plots,2) 
        libs.imshow(foldarr, extent=[mnn-0.5,size+mnn-0.5,-0.5,size-0.5])
        pl.xlabel('Fold'); pl.title('Residual')
      if maxsnr<thresh:
        pl.suptitle(title+' fold-start plot')
      else:
        pl.suptitle(title+' fold-start plot ('+str(maxs1)+' '+'%.1f' %(maxsnr))
      pl.savefig(savefig+'_foldstart.png')

    return str1, savefig+'_foldstart.png'

################################################################################################

def rms_ant(arr1, d):
    """
    arr is nscan X nchan for an ant-pol
    Calc median spectrum
    Subtract med spec from all scan specs, calc rms per chan and at gchan
    """
    import pylab as pl
    pl.ion()

    nscan, nchan = d.nscan, d.nchan
    if not hasattr(d, 'gchan'): 
      d.gchan = nchan/2
    gchan = d.gchan
    bchan, echan = d.bchan, d.echan
    arr = N.copy(arr1)

    if not d.gainuse:
        print("Gain calibrating before calculating median")
        # Calc 'gain' at gchan+/-8 chan
        dchan=8
        bp_gain = N.nanmedian(arr[:,gchan-dchan:gchan+dchan], 1)
        # Divide out the array by these 'gains'
        arr = N.asarray([arr[iscan]/bp_gain[iscan] for iscan in range(d.nscan)])

    # Calc median
    medsp = N.nanmedian(arr,0)  # over scan, is single spectrum

    # Subtract median sp from all sp. This gives scan X chan with mean 0 and no baseline
    arr = N.asarray([sub-medsp for sub in arr])

    # Stats over good chans over scan-chan
    mean, std = N.nanmean(arr[:,bchan:echan]), N.nanstd(arr[:,bchan:echan])
    med = N.nanmedian(arr[:,bchan:echan])
    mad = 1.5*N.nanmedian(N.abs(arr[:,bchan:echan]-med))

    # Calc std over scans for each good chan and get max val
    stddum = N.nanstd(arr[:,bchan:echan],0)
    maxstd = N.nanmax(stddum)

    # Stats over scans at gchan
    gmean, gstd = N.nanmean(arr[:,gchan]), N.nanstd(arr[:,gchan])
    gmed = N.nanmedian(arr[:,gchan])
    gmad = 1.5*N.nanmedian(N.abs(arr[:,gchan]-gmed))

    return medsp, N.asarray([mean, std, med, mad, maxstd]), N.asarray([gmean, gstd, gmed, gmad])

################################################################################################

def plot_med_ants(d):
    """
    Plot median spectra for each antenna for each pol
    And also zoom in to goodchans if defined
    medians is median spectra over (pol X ant x chan)
    """
    import pylab as pl
    pl.ion()

    nn = 1
    if hasattr(d, 'bchan') and d.echan-d.bchan < d.nchan:
        nn = 2; dogood = True

    # Plot the median spectrum for all ants for each pol and also for goodchans
    pl.figure()
    for ipol in range(d.npol):
      pl.subplot(nn,2,ipol+1)
      libs.imshow(N.transpose(d.bp_med[ipol]))
      pl.xlabel('Chan'); pl.ylabel('Antenna'); pl.title('Pol'+str(ipol))
      if dogood: 
        pl.xlabel(''); pl.xticks([])

      if dogood:
        pl.subplot(nn,2,ipol+1+2)
        libs.imshow(N.transpose(d.bp_med[ipol][:,d.bchan:d.echan]))
        pl.xlabel('Chan'); pl.ylabel('Antenna'); pl.title('Pol'+str(ipol)+' good')

    pl.suptitle(d.fname+' scan-median spectra')
    pl.savefig(d.pldir+d.fname+'_scanmed_spec.png') 

################################################################################################

def plot_rms_ants(d):
    """
    stats is for (pol X ant) and has (mean,std,med,mad,maxstd)
    gstats if it exists, is (pol X ant) and has (mean,std,med,mad) evaluated at gchan
    """
    import pylab as pl
    pl.ion()

    # Plot the statistics for all goodchans and gchans after subtracting median spec
    pl.figure(figsize=(7,4))
    for ipol in range(d.npol):
      pl.subplot(2,2,ipol+1)
      pl.plot(d.bp_stats[ipol,:,1],'ob'); pl.plot(d.bp_stats[ipol,:,3],'*g')
      #pl.plot(d.bp_stats[ipol,:,4],'.r')
      pl.title('Pol'+str(ipol))
      pl.subplot(2,2,ipol+3)
      pl.plot(d.bp_gstats[ipol,:,1],'ob'); pl.plot(d.bp_gstats[ipol,:,3],'*g')

    for i in range(4):
      pl.subplot(2,2,i+1)
      a1,a2,a3,a4=pl.axis()
      for i in range(d.nant/10):
          pl.plot([i*10, i*10],[a3,a4],'k-')

    pl.suptitle('Std(b), MAD(g), MaxStd-chan(r) for each ant BP (goodchans and gchan')
    pl.savefig(d.pldir+d.fname+'_scanmed_stats.png') 
    
################################################################################################

def calib_bp(d, doagain=False, doplot=True):
    """
    Calculate gains of bandpass table. But this can be done for autocorr too, as long
    as input is ant-based
    """
    import pylab as pl
    pl.ion()

    # Check if done already
    if hasattr(d, 'docalib') and (not doagain):
        print("Not doing calibration again. Set doagain to True and repeat")
    else:
        if not hasattr(d, 'dchan'): d.dchan = 5  # +/- to average around gchan
        d.bp_amp_gain = N.nanmean(d.bpamp[...,d.gchan-d.dchan:d.gchan+d.dchan],-1) 
                                                        # (npol X nscan X nant)
        d.bp_ph_gain = N.nanmean(d.bpph[...,d.gchan-d.dchan:d.gchan+d.dchan],-1)  
                                                        # Not dealing with wraps
        d.bp_gain_polrat = N.nanmedian(d.bp_amp_gain[0]/d.bp_amp_gain[1],  0)  # median ratio per ant
        d.docalib = True
    
        if doplot:
            pl.figure(figsize=(8,8))
            for i in range(d.nant):
                pl.subplot(6,10,i+1); pl.plot(d.bp_amp_gain[0,:,i]);
                pl.plot(d.bp_amp_gain[1,:,i]*d.bp_gain_polrat[i])
                pl.xticks([]); pl.title(str(i))
                if i>49: pl.xlabel('Scan')
                if i%10!=0: pl.yticks([])
            pl.suptitle('Bandpass gains (2pols) for each ant around chan'+str(d.gchan))
            pl.savefig(d.pldir+d.fname+'_bandpass_gains.png')
            
            pl.figure(figsize=(8,4))
            pl.subplot(121);pl.plot(d.bp_gain_polrat, '*-'); pl.xlabel('Ant'); pl.ylabel('Pol norm')
            pl.subplot(122);ret=pl.hist(d.bp_gain_polrat,20); 
            pl.suptitle('X/Y normalisation for bandpass gains')
            pl.savefig(d.pldir+d.fname+'_bandpass_gains_polnorm.png')

################################################################################################
        
def apply_calib_bp(d, doagain=False):
    """ Apply the gain calib calculated on BP on the BPs """

    if hasattr(d, 'gainuse') and (not doagain):
        print("Not applying calibration again. Set doagain to True if sure")
    else:
        if d.docalib:
            for ipol in range(d.npol):
                for iant in range(d.nant):
                    for iscan in range(d.nscan):
                        d.bpamp[ipol, iscan, iant, :] /= d.bp_amp_gain[ipol,iscan,iant]
            d.gainuse = True
        else:
            print("First run calib_bp to create gain table")
    
################################################################################################

def subtract_medsp(d):
    """ 
    Return median subtracted bp amps 
    d.med_sp is (polXantXchan) and d.bpamp is (polXscanXantXchan
    """

    d.bpamp_med = N.zeros(d.bpamp.shape)
    for ipol in range(d.npol):
        for iant in range(d.nant):
            for iscan in range(d.nscan):
                d.bpamp_med[ipol,iscan,iant] = d.bpamp[ipol,iscan,iant]-d.bp_med[ipol,iant]

    print("Created d.bpamp_med")

################################################################################################

def addflags_1(d, inputs):
    """
    addflags_1 has flags that apply for continuum (in principle, all chans)
    Each input is either 1 number, None or list
    All lists have to be same size. None is all and 1 number is just that
    """
 
    def getlist(i, nn, inp):
        dum = inp[i]
        if dum==None: return N.arange(nn[i])
        if isinstance(dum,int): return N.asarray([dum])
        if (isinstance(dum,list)) or (isinstance(dum,N.ndarray)): return dum

    if len(inputs) != 3:
        raise RuntimeError("Need 3 inputs for pols, scans, ants")

    nn = [d.npol, d.nscan, d.nant]
    inds = []
    for i in range(3):
      inds.append(N.asarray(getlist(i, nn, inputs)))

    flags = []
    for i in inds[0]:
        for j in inds[1]:
            for k in inds[2]:
                flags.append([i,j,k])

    if not hasattr(d, 'flags1'): d.flags1 = []
    if flags not in d.flags1:
        d.flags1.append(flags)
    else:
        print("Duplicate flag, not appending it")

################################################################################################

def addflags_2(d, inputs):
    """
    addflags_1 has flags that apply in spectral domain
    Each input is either 1 number, None or list
    All lists have to be same size. None is all and 1 number is just that
    But since this is for chans, if we write out for all pol,scan,ant, thats a lot
    So, write out only for pol and assume that all chans and ants are flagged
    That is, inputs are only for (pol, chans)
    """
 
    def getlist(i, nn, inp):
        dum = inp[i]
        if dum==None: return N.arange(nn[i])
        if isinstance(dum,int): return N.asarray([dum])
        if (isinstance(dum,list)) or (isinstance(dum,N.ndarray)): return dum

    if len(inputs) != 2:
        raise RuntimeError("Need 2 inputs for pols, chans")

    nn = [d.npol, d.echan-d.bchan]
    inds = []
    for i in range(2):
      inds.append(N.asarray(getlist(i, nn, inputs)))

    flags = []
    for i in inds[0]:
       for j in inds[1]:
           flags.append([i,j])

    if not hasattr(d, 'flags2'): d.flags2 = []
    if flags not in d.flags2:
        d.flags2.append(flags)
    else:
        print("Duplicate flag, not appending it")

################################################################################################

def chan2ripple(arr, ind0, bchan=0, echan=None, domf=False, win=31):
    """
    See if there is a significant 2 channel ripple in index ind0
    If ndim=4 then assume it is pol X (1,2,ind0 random order)
    If ndim=3 then assume it is pol X (1,ind0 random order)
    Ripple axis is from bchan to echan
    """
    if echan==None: echan = arr.shape[ind0]

    if arr.ndim==4:
        if arr.shape[0] != 2:
            raise RuntimeError("Need first axis to be npol=2")
        myarr = N.moveaxis(arr,ind0,-1)  # move ripple axis to end
        npol, n1, n2, nrip = myarr.shape
        ringind = N.zeros((npol,n1,n2))
        for ipol in range(npol):
            for in1 in range(n1):
                for in2 in range(n2):
                    x = myarr[ipol,in1,in2,bchan:echan]
                    if domf:
                        x = x - poly_filter(x, win)
                    xroll = N.roll(x, 1)
                    dum = N.where(xroll<x,-1,1)
                    if N.sum(N.isnan(x))>0:
                        inds = N.where(N.isnan(x)); dum = dum*1.0
                        dum[inds] = N.nan
                        if inds[-1][-1]+1==len(dum): inds[-1][-1]=0
                        dum[inds[-1]+1] = N.nan
                    l1 = len(dum[::2]); l2 = len(dum[1::2])
                    ringind[ipol,in1,in2] = len(N.where(dum[::2][:l2]+dum[1::2]==0)[0])
                    nn = len(x) - N.sum(N.isnan(x))
                    ringind[ipol,in1,in2] = ringind[ipol,in1,in2]*6.0/nn-2.0

    if arr.ndim==3:
        if arr.shape[0] != 2:
            raise RuntimeError("Need first axis to be npol=2")
        myarr = N.moveaxis(arr,ind0,-1)  # move ripple axis to end
        npol, n1, nrip = myarr.shape
        ringind = N.ones((npol,n1))
        for ipol in range(npol):
            for in1 in range(n1):
                x = myarr[ipol,in1,bchan:echan]
                if domf:
                    x = x - poly_filter(x, win)
                xroll = N.roll(x, 1)
                dum = N.where(xroll<x,-1,1)
                if N.sum(N.isnan(x))>0:
                    inds = N.where(N.isnan(x)); dum = dum*1.0
                    dum[inds] = N.nan 
                    if inds[-1][-1]+1==len(dum): inds[-1][-1]=0
                    dum[inds[-1]+1] = N.nan
                dum = dum[1:-2]
                l1 = len(dum[::2]); l2 = len(dum[1::2])
                ringind[ipol,in1] = len(N.where(dum[::2][:l2]+dum[1::2]==0)[0])
                nn = len(x) - N.sum(N.isnan(x))
                ringind[ipol,in1] = ringind[ipol,in1]*6.0/nn-2.0

    if arr.ndim==1:
        x = arr[bchan:echan]
        if domf:
            x = x - poly_filter(x, win)
        xroll = N.roll(x, 1)
        dum = N.where(xroll<x,-1,1)
        if N.sum(N.isnan(x))>0:
            inds = N.where(N.isnan(x)); dum = dum*1.0
            dum[inds] = N.nan 
            if inds[-1][-1]+1==len(dum): inds[-1][-1]=0
            dum[inds[-1]+1] = N.nan
        l1 = len(dum[::2]); l2 = len(dum[1::2])
        ringind = len(N.where(dum[::2][:l2]+dum[1::2]==0)[0])
        nn = len(x) - N.sum(N.isnan(x))
        ringind = ringind*6.0/nn-2.0

    return ringind

################################################################################################

def write_file(sname, var, nsave, verbose):

    if nsave:
        N.save(sname, var)
    else:
        pickle.dump(var, open(sname+'.pickle',"w"))
    if verbose: print("   Wrote out", sname)

################################################################################################

def oorgai(d, f, nsave, scans, bchan, echan, npol=2, getflags=True, dotimeavg=True, \
           dofullauto=True, onlyauto=False, dofullcross=False, verbose=False):
    """
    Read vis and flags for a scan and write out pickled files for 
    cross and auto, time avg with flags and full oens without, and also dc
    Note that time avg are float and without they are int
    """
    from sys import getsizeof

    t1 = time.time(); times1, times2, times3 = [], [], []
    f.select()
    ch = range(bchan,echan)

    for iscan in scans:
     #try:
        t11 = time.time()
        print("Working on scan", iscan)
        tt1 = time.time()

        scandir = 'scan_'+str(iscan)
        odir = d.sdpdir+ddir(scandir)
        if not os.path.isdir(odir): os.mkdir(odir)
        sname = odir+'/scan_'+str(iscan)
        
        if onlyauto:
          print("Reading only autocorr")
          f.select(scans=iscan, channels=ch)
          if f.size/1e9 > 300.0:
            print("Large file, reading auto by each dump")
            f.select(scans=iscan, corrprods='auto', pol='hh,vv', channels=ch)
            vis = []
            for idump in len(f.dumps):
              vis.append(f.vis[idump])
              vis = N.asarray(vis)
          else:
            f.select(scans=iscan, corrprods='auto', pol='hh,vv', channels=ch)
            vis = f.vis[:]
        else:
          f.select(scans=iscan, pol='hh,vv', channels=ch)
          vis = f.vis[:]
        if getflags:
          flag = f.flags[:]
          flag = N.where(flag==1, N.nan, 1)
        else: flag=None
        tt2 = time.time(); times1.append(tt2-tt1)
        print("   Shape of vis is", vis.shape, str(getsizeof(vis)/1e9)[:5], 'GB')
        print("   Done reading vis, flags for scan", iscan)

        if dotimeavg:
          tt1 = time.time()
          vis_avt = timeavg_comp(vis, flag, d.nant, d.nbl, npol, echan-bchan) # returns t-avged data
          tt2 = time.time(); times2.append(tt2-tt1)
          print("   Done computing stats for scan", iscan)
        
          tt1 = time.time()
          write_file(sname+'_cross_scanav', vis_avt[0], nsave, verbose)
          write_file(sname+'_auto_scanav', vis_avt[1], nsave, verbose)
          del vis_avt

        # Pickle out the full auto and vis data if asked for
        if dofullcross:
          fullcross = []
          for i in range(npol):
            fullcross.append(vis[:,:,2*d.na+i*d.nbl:2*d.na+(i+1)*d.nbl])
          write_file(sname+'_fullcross', N.asarray(fullcross), nsave, verbose)
          write_file(sname+'_fullcrossdc', N.asarray(fullcross)[:,:,0,:], nsave, verbose)
          del fullcross

        if getflags:
          fullflags = []
          for i in range(npol):
              fullflags.append(flag[:,:,2*d.na+i*d.nbl:2*d.na+(i+1)*d.nbl])
          write_file(sname+'_fullcrossflags', N.asarray(fullflags), nsave, verbose)
          del fullflags
           
        if dofullauto:
          fullauto = []
          for i in range(npol):
            fullauto.append(vis[:,:,i*d.na:(i+1)*d.na])
          write_file(sname+'_fullauto', N.abs(N.asarray(fullauto)), nsave, verbose)
          write_file(sname+'_fullautodc', N.abs(N.asarray(fullauto))[:,:,0,:], nsave, verbose)
          del fullauto

        if getflags:
          fullflags = []
          for i in range(npol):
            fullflags.append(flag[:,:,i*d.na:(i+1)*d.na])
          write_file(sname+'_fullautoflags', N.asarray(fullflags), nsave, verbose)
          del fullflags

        del vis, flag
        tt2 = time.time(); times3.append(tt2-tt1)
        print("   Done writing out scan", iscan, "in ", odir)
        
        t21 = time.time()
        print("Done scan", iscan, "in ", (t21-t11)/60., "min")
        print()
     #except:
     #  print("***************************")
     #  print("**FAILED ON SCAN",iscan,"****")
     #  print("***************************")
    t2 = time.time()
    print("%s %.1f %s" %("Reading all scans in", (t2-t1)/60, "min"))
    print("%s %.1f %s" %("Time for reading in each scan", N.mean(N.asarray(times1))/60, "min"))
    print("%s %.1f %s" %("Time for averaging each scan ", N.mean(N.asarray(times2))/60, "min"))
    print("%s %.1f %s" %("Time for pickling each scan  ", N.mean(N.asarray(times3))/60, "min"))
    
    print("Wrote into ", d.sdpdir)

################################################################################################

def pickle_new_data(num, nsave=True, getflags=True, dotimeavg=True, dofullauto=True, \
    dofullcross=True, onlyauto=False, verbose=True, extn=None, basedir='/data/mohan/', bchan=0, echan=0):
    """
Read in a rdb file and writes our scan based files. Writes out numpy files by default. Writes out only
XX and YY data for now.

num = CBID
nsave=True (default) writes out numpy files, else pickle files
getflags=True (default) will read in flags and 1. write out data after applying them and 2. write out flags
dotimeavg=True (default) will write out scan averaged auto and visibility data files separately
dofullauto=True (default) will write out the full auto corr data for each scan as well
dofullcross=True (default) will write out the full visibility data for each scan as well
onlyauto=False (default) will write out visibility data as well
verbose=True (default) will print out more info
extn=None (default) will assume the full rdb file. For simple file use extn='none'
basedir is the directory where the CBID folder is (default is /data/mohan)
bchan, echan are the channel range to write out data (default is full range)
    """

    import katdal
    class myclass(object):
        pass

    basedir += '/'
    
    d = myclass()

    if onlyauto and (dofullcross or dotimeavg or getflags):
      raise RuntimeError("Onlyauto can be True only if dotimeavg, getflags and dofullcross are False")
    
    fname = num
    create_dirs(d, fname, basedir)

    flog = open(d.sdpdir+'/pickle.log','w')
    
    fname = d.fname
    pldir = d.pldir
    
    t1 = time.time()
    str1 = d.fdir+d.fname+'_sdp_l0.'
    if extn==None: 
        str1 += 'full.rdb'
    else: 
        if extn=='none':
            str1 += 'rdb'
        else:
            str1 += extn + '.full.rdb'
    f = katdal.open(str1)
    t2 = time.time()
    print("Done opening", d.fdir+d.fname, " in ", t2-t1, " secs")
    flog.write("Done opening "+ d.fdir+d.fname+ " in "+ str(t2-t1)+ " secs\n")
    
    read_katdal_para(f, d, flog=flog)
    if echan==0: echan=d.nchan
    print("Taking channels", bchan, "to", echan)
    
    f.select(scans='track')
    print(f)
    
    d.gscans = f.scan_indices
    
    print("Good scans are")
    sources = []
    f.select()
    for g in d.gscans: 
        f.select(scans=g)
        print(g, f.vis.shape)
    
    ret = input('Take these ? (y/n) ')
    #ret = 'y'
    if ret=='n':
      dum = []
      while True:
        ret = input('')
        if ret=='n': break
        dum.append(int(ret))
        d.gscans = N.asarray(dum)
    print("Taking ", d.gscans)
    
    d.npol = npol = 2
    
    os.system('date')
    # Pickle it all out
    try:
        oorgai(d, f, nsave, d.gscans, bchan, echan, npol, getflags, dotimeavg, dofullauto, onlyauto, \
                      dofullcross, verbose)
    except:
        os.system('Failed on date')


################################################################################################

def write_antnames(num, extn):

    import katdal

    fdir = '/data/mohan/'
    fdir = ddir(fdir)

    # Get ant names

    if extn=='': 
      fname = fdir+num+'/'+num+'_sdp_l0.full.rdb'
    else:
      if extn=='none':
        fname = fdir+num+'/'+num+'_sdp_l0.rdb'
      else:
        fname = fdir+num+'/'+num+'_sdp_l0.full.rdb'

    if not os.path.isfile(fname): raise RuntimeError(fname+' not found')
    f = katdal.open(fname)
    nant = len(f.ants)
    names = []
    for iant in range(nant):
        names.append(f.ants[iant].name)
    names = N.asarray(names)

    fn = open(fdir+num+'/'+num+'_pickle/antnames',"w")
    for name in names:
        fn.write(name+'\n')
    fn.close()

################################################################################################

def write_startstop(num):

    import katdal

    fdir = '/data/mohan/'
    fdir = ddir(fdir)

    # Get ant names
    num = str(num)
    fname = fdir+num+'/'+num+'_sdp_l0.full.rdb'
    if not os.path.isfile(fname): 
      fname = fdir+num+'/'+num+'_sdp_l0.rdb'
      if not os.path.isfile(fname): 
        raise RuntimeError(fname+' not found')
    print("Reading", fname)
    f = katdal.open(fname)
    nscan = len(f.scan_indices)
    print("Number of scans", nscan)

    beg, end = [], []
    dump = f.dump_period
    for i in range(nscan):
      f.select(scans=i)
      beg.append(f.timestamps[0])
      end.append(f.timestamps[-1])
    beg, end = N.asarray(beg), N.asarray(end)

    N.save(fdir+num+'/'+num+'_pickle/scantimes', [dump, beg, end])
    print("Wrote", fdir+num+'/'+num+'_pickle/scantimes.npy')

################################################################################################

def sep_auto_ant(num, scans=None, extn=''):
    """
    Get all auto and rewrite separate files for each antenna
    Scans can be list of list of scans to group by. Blank times are Nan-ed.
    Assume that the scans numbers within a group are time ordered.
  """
    import glob, pickle

    fdir = '/data/mohan/'
    fdir = ddir(fdir)

    if isinstance(num,int): num = str(num)

    # read ant names
    fname = fdir+num+'/'+num+'_pickle/antnames'
    if not os.path.isfile(fname): 
        print("Running write_antnames")
        write_antnames(num, extn)
    fn = open(fname)
    names = []
    for line in fn:
        names.append(line.strip())
    fn.close()
    names = N.asarray(names)

    # Get scan dir names
    ffs = glob.glob(fdir+num+'/'+num+'_pickle/scan_*')
    ffs = N.asarray(ffs)
    nscan = len(ffs)
    if nscan==0: raise RuntimeError("No scans found")

    # Sort scan dir names
    dum = []
    for f in ffs:
        dum.append(int(f.split('_')[-1]))
    dum = N.asarray(dum)
    ffs[N.argsort(dum)]
    ffs = ffs[N.argsort(dum)]

    # Get beg and end times to add Nans between scans if needed
    fname = fdir+num+'/'+num+'_pickle/scantimes.npy'
    if not os.path.isfile(fname): 
        print("Running write_startstop")
        write_startstop(num)
    times = N.load(fdir+num+'/'+num+'_pickle/scantimes.npy', allow_pickle=True)
    dump, begs, ends = times

    # Groups of scans
    if scans==None: 
        ngroup = 1
        scans = [N.sort(dum)]
    else:
        ngroup = len(scans)

    for ii in range(ngroup):
        print('Group', ii+1)
        for jj,iscan in enumerate(scans[ii]):
            print("  Opening scan_"+str(iscan), )
            fname = fdir+num+'/'+num+'_pickle/scan_'+str(iscan)+'/scan_'+str(iscan)+'_fullauto'
            data = myload(fname)
            print(data.shape)
            
            if jj==0: 
                arr = data
            else:
                dt = int(round((begs[jj]-ends[jj-1])/dump))
                print('  Adding',dt,'dumps')
                a1,a2,a3,a4 = data.shape
                dummy = N.zeros((a1,dt,a3,a4))*N.nan
                arr = N.concatenate((arr,dummy),1)
                arr = N.concatenate((arr,data),1)

        print("  Shape of final array is", arr.shape)
        dname = fdir+num+'/'+num+'_pickle/autoants/'
        if not os.path.isdir(dname): os.mkdir(dname)

        for iant in range(arr.shape[3]):
            fname = dname+'/auto_group'+str(ii)+'_ant'+str(iant)+'_'+names[iant]
            N.save(fname, arr[:,:,:,iant])
            print("  Wrote", fname+'.npy')

################################################################################################

def dedisp(arr, docalib=True, bx=None,ex=None,by=None,ey=None, flagy=None, flatten=True, \
           boxsize=50, thresh1=3.0, threshneg=-4.0, maxdm=100, minsnr=20, axis=1, niter=3,\
           doabs=False, robust=False, peaksnr=10.0, sidesnr=3.0, dmdiff=3, doflag=True, doplot=False):

    if arr.ndim!=2: raise RuntimeError("Array needs to be 2-dimensional")
    nx, ny = arr.shape

    if bx==None: bx=0
    if by==None: by=0
    if ex==None: ex=nx
    if ey==None: ey=ny
    data = N.copy(arr[bx:ex,by:ey])
    if len(flagy)>0:
        for flag in flagy:
            data[:,flag[0]-by:flag[1]+1-by] = N.nan
    inputdata = N.copy(data)
 
    # flag narrow band rfi left over after applying my mask. So use nanmax instead of nanmean
    if doflag:
        bp0 = N.nanmax(data,0)
        bp = bp0 - poly_filter(bp0, 51)
        med, mad, dum = getmad(bp, niter=3, thresh=3.0, doabs=False, robust=False)
        #print('stats ',med, mad)

        inds = N.where(N.abs((bp-med)/mad)>5.0)[0]
        if len(inds)>0:
            for ind in inds: 
                data[:,ind-2:ind+3] = N.nan
                bp[ind-2:ind+3] = N.nan

        """
        import pylab as pl
        pl.ion()
        pl.figure()
        pl.subplot(311); pl.plot(bp0)
        pl.subplot(312); pl.plot(bp)
        pl.subplot(313); libs.imshow(data)
        #return None
        """


    # Calibrate on both axis
    if docalib:
        bp = N.nanmean(data,0)
        for i in range(data.shape[0]):
            data[i,:] /= bp

        lc = N.nanmean(data,1)
        for i in range(data.shape[1]):
            data[:,i] /= lc

    """
    pl.figure()
    pl.subplot(221); libs.imshow(data)
    pl.subplot(222); pl.plot(N.nanmax(data,1))
    pl.subplot(223); pl.plot(N.arange(len(N.nanmax(data,0)))+4200,N.nanmax(data,0))
    pl.subplot(224); pl.plot(N.arange(len(N.nanmax(data,0)))+4200,N.nanmean(data,0))
    #return None
    """

    # Flatten by removing 2d polynomial over boxsizeXboxsize boxes
    if flatten:
        dum = N.copy(data)

        # first change Nan (time breaks) to rms for linalg then nan then after
        nans = N.where(N.isnan(N.nanmean(dum,1)))[0]
        if len(nans)>0:
            good = N.where(~N.isnan(N.nanmean(dum,1)))[0]
            for nan in nans:
                # find nearest non-nan neighbours on both sides
                ii = N.searchsorted(good,nan)
                if ii<len(good):
                  l1, l2 = good[ii-1], good[ii]
                  data[nan] = 0.5*(dum[l1]+dum[l2])
                else:
                  l1 = good[ii-1]
                  data[nan] = dum[l1]
        # first change Nan (channel breaks) to rms for linalg then nan then after
        nans = N.where(N.isnan(N.nanmean(dum,0)))[0]
        if len(nans)>0:
            good = N.where(~N.isnan(N.nanmean(dum,0)))[0]
            for nan in nans:
                # find nearest non-nan neighbours on both sides
                ii = N.searchsorted(good,nan)
                if ii<len(good):
                  l1, l2 = good[ii-1], good[ii]
                  data[:,nan] = 0.5*(dum[:,l1]+dum[:,l2])
                else:
                  l1 = good[ii-1]
                  data[:,nan] = dum[:,l1]

        xr = range(0,data.shape[0],boxsize)
        if xr[-1]+boxsize-data.shape[0] > boxsize/2:
            xr = range(0,data.shape[0]-boxsize,boxsize)

        yr = range(0,data.shape[1],boxsize)
        if yr[-1]+boxsize-data.shape[1] > boxsize/2:
            yr = range(0,data.shape[1]-boxsize,boxsize)

        for i in xr:
            if i==xr[-1]: end1 = data.shape[0]
            else: end1 = i+boxsize

            for j in yr:
                if j==yr[-1]: end2 = data.shape[1]
                else: end2 = j+boxsize

                z = N.copy(data[i:end1,j:end2])

                med = N.nanmedian(z); mad = 1.5*N.nanmedian(N.abs(z-med))
                inds = N.where(N.abs(z-med)/mad>thresh1)
                z[inds] = N.random.normal(med,mad,z[inds].shape)

                x,y = N.mgrid[0:z.shape[0],0:z.shape[1]]
                X, Y = x.flatten(), y.flatten()
                A = N.array([X*0+1, X, Y, X*X, Y*Y, X*Y]).T
                B = z.flatten()
                c, r, rank, s = N.linalg.lstsq(A, B)
                z1 = c[0] + c[1]*x + c[2]*y + c[3]*x*x + c[4]*y*y + c[5]*x*y
                dum[i:end1,j:end2] -= z1

    flat = N.copy(dum)
    # identify broadband RFI which will slip through if it has -ve and +ve in dum
    lc = N.nanmean(N.where(dum>0,dum,N.nan),1)  # ignore negatives
    mean, std, dum1 = getmad(lc, niter=3, doabs=False, robust=False)
    inds = N.where((lc-mean)/std>=5.0)[0]
    for ind in inds:
        dum[ind] = N.nan
    flat1 = N.copy(dum)

    # Put all dips as 1 and Nan everything else
    med = N.nanmedian(dum); mad = 1.5*N.nanmedian(N.abs(dum-med))
    inds = N.where((dum-med)/mad > threshneg)
    dum[inds] = N.nan
    dum = N.abs(dum)
    dum = dum/dum
    normed = N.copy(dum)
                
    dum2 = []
    for roll in range(-maxdm,maxdm):
        dum3 = N.copy(dum)
        for i in range(dum.shape[1]):
            disp = int(round(i*roll*1.0/dum.shape[1]))
            dum3[:,i] = N.roll(dum[:,i],disp,0)
        av = N.nansum(dum3,1)
        dum2.append(av)
    dum2 = N.transpose(N.asarray(dum2))
    dum = N.copy(dum2)

    if doplot:
        import pylab as pl
        pl.ion()
        pl.figure()
        pl.subplot(231);libs.imshow(inputdata); pl.title('Input data')
        pl.subplot(232);libs.imshow(flat); pl.title('Detrended')
        pl.subplot(233);libs.imshow(flat1); pl.title('Flagged data data')
        pl.subplot(234);libs.imshow(normed); pl.title('Segmented')
        pl.subplot(235);libs.imshow(dum); pl.title('Radon transform')


    # dum2 is the radon transform. Now to detect pulses and 'associate' them
    store = []
    med0, mad0, dum = getmad(dum2, axis=axis, niter=niter, doabs=doabs, robust=False)

    
    while True:
        mx0 = N.unravel_index(N.nanargmax(dum2),dum2.shape)
        med,mad = med0[mx0[0]], mad0[mx0[0]]
        if mad==0: mad = 1.0
        if ((dum2[mx0]-med)/mad>peaksnr) and dum2[mx0]>minsnr:
            y = (dum2[mx0[0]]-med)/mad
            store.append([mx0[0],mx0[1]-maxdm,(dum2[mx0]-med)/mad])

            # This is sidesnr sigma on both sides of dm of max
            l2 = N.where(y[mx0[1]:]<sidesnr)[0][0]+mx0[1]-1
            l1 = N.where(y[:mx0[1]]<sidesnr)[0][-1]+1
            dum2[mx0[0],l1:l2+1] = N.nan
            
            lim = [dum2.shape[0],0]; sgn = [1,-1]
            for ii in range(2): # each side
                newmx = mx0[1]  # dm of max, to be updated as we move to side
                for side in range(mx0[0]+sgn[ii],lim[ii],sgn[ii]):
                    thismx = N.argmax(dum2[side])
                    if abs(newmx-thismx)<=dmdiff and abs(thismx-mx0[1])<20:  # associated pulse
                        med,mad = med0[side], mad0[side] 
                        y1 = (dum2[side]-med)/mad

                        l2 = N.where(y1[thismx:]<sidesnr)[0][0]+thismx-1
                        dum4 = N.where(y1[:thismx]<sidesnr)[0]
                        if len(dum4)>0:
                            l1 = N.where(y1[:thismx]<sidesnr)[0][-1]+1
                        else:
                            l1 = 0
                        dum2[side,l1:l2+1] = N.nan
                        newmx = thismx                    
                    else:
                        break
        else:
            break

    store = N.asarray(store)
    return data.shape, store

################################################################################################

def get_rfi(num, extn=None, archive=False, prod='auto'):
    import katdal, glob 
    import corr_test_chans as cc
    import pylab as pl
    pl.ion()

    basedir = '/scratch2/mohan/'
    basedir += '/'
    if isinstance(num,int): num = str(num)

    if archive:
      if extn==None:
          fn = basedir+num+'/'+num+'_sdp_l0.full.rdb'
      else:
        if extn=='none':
          fn = basedir+num+'/'+num+'_sdp_l0.rdb'
        else:
          fn = basedir+num+'/'+num+'_sdp_l0.'+extn+'.full.rdb'
      print("Opening ", fn)
      f = katdal.open(fn)

      fr = f.channel_freqs
      N.save(open(basedir+'/'+num+'/freqs.npy','w'), fr)
      print("Wrote", basedir+'/'+num+'/freqs.npy')
      
      f.select(corrprods=prod, pol='hh,vv', scans='track')
      data = N.abs(f.vis[:])
      nt, nchan, ncorr = data.shape
      print("Shape of auto corrs", data.shape)
    else:  # from disk
      fdir = basedir+num+'/'+num+'_pickle/'
      fns = N.sort(glob.glob(fdir+'scan_*'))
      data = []
      for fn in fns:
        data.append(N.nanmean(N.load(open(fn+'/'+fn.split('/')[-1]+'_'+prod+'_scanav.npy'))[0],0))  # npol nchan ncorr
      data = N.asarray(data)
      fr = N.load(open(basedir+'/'+num+'/freqs.npy'))
      nt, nchan, ncorr = data.shape
      print("Shape of auto corrs", data.shape)

    avspec = []
    antspecs = N.zeros((nchan,ncorr))
    for iant in range(ncorr):
        x = N.nanmean(N.abs(data[:,:,iant]),0)*1.0
        antspecs[:,iant] = x
        x = x/N.nanmedian(x[len(x)/2-50:len(x)/2+50])
        avspec.append(x)
    avspec = N.nanmean(avspec,0)

    maxspec = N.zeros(nchan)
    for ichan in range(nchan):
        x = antspecs[ichan]
        x = N.where(N.isnan(x),-99,x)
        maxspec[ichan] = N.sort(x)[-3]  # cos max may be crap ant
    #maxspec = maxspec/N.nanmedian(maxspec[len(maxspec)/2-50:len(maxspec)/2+50])

    l1 = nchan/20; l2 = nchan-l1
    pl.figure(figsize=(10,8))
    pl.subplot(221)
    pl.semilogy(N.arange(l1,l2),avspec[l1:l2])
    pl.xlabel('Channel')
    pl.subplot(222)
    pl.semilogy(fr[l1:l2]/1e6, avspec[l1:l2])
    pl.xlabel('Frequency (MHz)')

    mybands = cc.mybands['l'][32]
    mybandflags = cc.mybandflags['l'][32]
    myfullband = cc.myfullband['l'][32]
    myfullbandflags = cc.myfullbandflags['l'][32]

    for ib in range(len(mybandflags)):
      for ibb in range(len(mybandflags[ib])):
        bb = mybandflags[ib][ibb]
        avspec[bb[0]:bb[1]] = N.nan
        maxspec[bb[0]:bb[1]] = N.nan

    pl.subplot(223)
    for iband in range(len(mybands)):
      pl.semilogy(N.arange(l1,l2)[mybands[iband][0]-l1:mybands[iband][1]-l1], \
                    avspec[l1:l2][mybands[iband][0]-l1:mybands[iband][1]-l1])
    pl.xlabel('Channel')
    pl.subplot(224)
    pl.semilogy(fr[l1:l2]/1e6, avspec[l1:l2])
    pl.xlabel('Frequency (MHz)')
    
    pl.suptitle(str(num)+' '+extn+' average spectrum vs chan, freq for '+prod)
    pl.savefig(str(num)+'_'+extn+'_avspec_'+prod+'.png')
    
################################################################################################

def get_64chandip(fn):
  
  import os, sys, glob,  warnings
  import corr_test_chans as cc
  import numpy as N
  warnings.filterwarnings("ignore")

  fold = 256
  clip = 0.5
  # read data
  data = N.load(fn)[0] # mean

  # cut out my band, apply flags, get amplitude
  fband = cc.myfullband['l'][32]
  fbandfl = cc.myfullbandflags['l'][32]
  data = data[:,fband[0]:fband[1],:]
  for ifl in range(len(fbandfl)):
    data[:,fbandfl[ifl][0]-fband[0]:fbandfl[ifl][1]-fband[0],:] = N.nan
  amp = N.abs(data)
  del data
  inds = N.where(amp<clip)
  amp[inds] = N.nan

  namp = N.copy(amp)
  for i in range(amp.shape[0]):
    for j in range(amp.shape[2]):
      namp[i,:,j] /= N.nanmedian(amp[i,:,j])
  normfullavg = N.nanmean(N.nanmean(namp,2),0)
  normfullavg_filt = normfullavg - poly_filter(normfullavg, 100)
  normfullavg_fold = N.asarray([N.nanmean(normfullavg_filt[i::fold]) for i in range(fold)])
  x = normfullavg_fold[N.where(~N.isnan(normfullavg_fold))[0]]
  dip = abs(N.nanmean(N.sort(x)[:fold/64]))

  first = 64 - fband[0]%64
  dipexact = abs(N.nanmean(x[first::64]))
 
  return dip, dipexact

################################################################################################

def fold(x, fold):

    return N.asarray([N.nanmean(x[k::fold]) for k in range(fold)])

################################################################################################

def get_rfi_chans(freq0, bw, nchan, drop=0.04):
  from myrfimask import myrfimask as myrfi

  flagchans = []
  for i in range(len(myrfi)):
    x,y = myrfi[i]
    x,y = x*1e6, y*1e6
    ch1, ch2 = int((x-freq0)*nchan/bw), int((y-freq0)*nchan/bw)
    if ch1>0 and ch2>0 and ch1<nchan and ch2<nchan:
      flagchans.append([ch1,ch2])
  flagchans.append([0,int(drop*nchan)])
  flagchans.append([int((1.-drop)*nchan),nchan])

  return flagchans
  
################################################################################################
    

def read_gains(caltable):
    """ 
    Read calibration tables (gaincal, bandpass, delaycal) written out by CASA
    Returns data, antpos, names, err, snr, flag, u, v, (index0,index1,antnums,antind), [scans, nant, field, nchan, freq0, bw]
    """

    import casacore.tables as t
    import glob, math

    def polfirst(arr):
        dum = []
        for i in range(arr.shape[-1]): # npol
            dum.append(arr[...,i])
        return N.asarray(dum)

    tt = t.table(caltable)

    times = tt.getcol('TIME')
    ant1 = tt.getcol('ANTENNA1')
    ant2 = tt.getcol('ANTENNA2')
    scans = tt.getcol('SCAN_NUMBER')
    if caltable[:5]=='delay':
      data = tt.getcol('FPARAM')
    else:
      data = tt.getcol('CPARAM')
    err = tt.getcol('PARAMERR')
    flag = tt.getcol('FLAG')
    snr = tt.getcol('SNR')
    field = tt.getcol('FIELD_ID')

    tt1 = t.table(caltable+'/ANTENNA')
    names = N.asarray(tt1.getcol('NAME'))
    antpos = N.asarray(tt1.getcol('POSITION'))
    #N.save('antposfull',antpos)
    for i in range(3): antpos[:,i] -= N.mean(antpos[:,i])

    tt2 = t.table(caltable+'/SPECTRAL_WINDOW/')
    dum3 = tt2.getcol('CHAN_FREQ')
    if len(dum3[0])==1:
      freq0= freq1 = dum3[0]
    else:
      freq0 , freq1 = dum3[0][0], dum3[0][1]

    if caltable[:5]=='delay':
      data = data*2*math.pi*1e-9*freq0*1.5

    ncorr, nchan, npol = data.shape
    ntimes = len(N.unique(times))
    #nscan = len(N.unique(scannum))
    nant = len(N.unique(ant1))
    ants = N.unique(ant1)
    bw = (freq1-freq0)*nchan
    
    if data.shape[0] == ncorr: print("Reshaping data array")

    print('ntimes, nant, nchan, npol = ', ntimes, nant, nchan, npol)

    if data.shape[0] == ncorr:
      data = data.reshape((ntimes, nant, nchan, npol))
      err = err.reshape((ntimes, nant, nchan, npol))
      snr = snr.reshape((ntimes, nant, nchan, npol))
      flag = flag.reshape((ntimes, nant, nchan, npol))
      field = field.reshape((ntimes, nant))[:,0]
      flag = N.where(flag,N.nan,1)

      data = polfirst(data)
      err = polfirst(err)
      snr = polfirst(snr)
      flag = polfirst(flag)

    fns = glob.glob('1*sdp_l0.rdb')
    if len(fns)==1:
      print("Reading u,v from",fns[0])
      import katdal
      from katsdpcal.calprocs import get_bls_lookup
      f = katdal.open(fns[0])
      f.select(scans='track',corrprods='cross', pol='h')
      u, v = f.u, f.v
      u = u[int(u.shape[0]/2),:]
      v = v[int(v.shape[0]/2),:]
      
      na = len(f.ants); nbl = int(na*(na-1)/2)
      ant_names = [a.name for a in f.ants]
      cross_blslook = get_bls_lookup(ant_names, f.corr_products)[:nbl]
      index0=[cp[0] for cp in cross_blslook]
      index1=[cp[1] for cp in cross_blslook]
      index0, index1 = N.asarray(index0), N.asarray(index1)
      antnums = [int(a[1:]) for a in ant_names]
      antind = N.asarray(N.ones((na,na))*999,int)
      for i in range(nbl):
        antind[index0[i],index1[i]] = i
        antind[index1[i],index0[i]] = i
    else:
      u, v = N.nan, N.nan

    return data, antpos, names, err, snr, flag, u, v, (index0,index1,antnums,antind), [scans, nant, field, nchan, freq0, bw]

################################################################################################
def plot_gain_bandpass(caltable, chan, num=None, doplot=True):

    import matplotlib 
    from myrfimask import myrfimask as myrfi
    matplotlib.use("Agg")
    import pylab as pl
    pl.ion()

    data, antpos, names, err, snr, flag, u, v, index, dum = read_gains(caltable)
    if isinstance(chan,int):
      data = data[:,:,:,chan]
      err = err[:,:,:,chan]
      snr = snr[:,:,:,chan]
      flag = flag[:,:,:,chan]
    else:
      data = N.nanmean(data[:,:,:,chan[0]:chan[1]],3)
      err = N.nanmean(err[:,:,:,chan[0]:chan[1]],3)
      snr = N.nanmean(snr[:,:,:,chan[0]:chan[1]],3)
      flag = N.nanmean(flag[:,:,:,chan[0]:chan[1]],3)

    a1,a2,a3,a4,a5,a6 = dum
    dum = a1, a2, a3, 1, a5, a6
    calsoln = data, antpos, names, err, snr, flag, u, v, index, dum

    plot_gaincal(calsoln, calname='sn_'+caltable+'_ch'+str(chan), num=num)
    
################################################################################################
def plot_twogains_bandpass(caltable1, caltable2, chan1, chan2, title):

    import matplotlib
    from myrfimask import myrfimask as myrfi
    matplotlib.use("Agg")
    import pylab as pl
    pl.ion()

    phs = []
    chans = [chan1, chan2]
    for ii,caltable in enumerate([caltable1, caltable2]):
      data, antpos, names, err, snr, flag, u, v, index, dum = read_gains(caltable)
      if isinstance(chans[ii],int):
        data = data[:,:,:,chans[ii]]
        flag = flag[:,:,:,chans[ii]]
      else:
        data = N.nanmean(data[:,:,:,chans[ii][0]:chans[ii][1]],3)
        flag = N.nanmean(flag[:,:,:,chans[ii][0]:chans[ii][1]],3)
      data = data*flag
      npol, nscan, nant = data.shape
      ph = N.angle(data,1)
      med = N.nanmedian(ph, 1)
      ph = ph - N.repeat(N.expand_dims(med,1),nscan,1)
      ph = N.where(ph>180, ph-360, ph)
      ph = N.where(ph<-180, ph+360, ph)
      phs.append(ph)
    phs = N.asarray(phs)

    for ipol in range(2): 
      pl.figure(figsize=(24,20))
      for iph in range(2):
        nn, mm = libs.subplot(nant-1)
        mn, mx = 1e7, -1e7
        for iii in range(len(phs[iph][ipol])): 
          mn,mx = min(mn, N.nanmin(phs[iph][ipol][iii])), max(mx, N.nanmax(phs[iph][ipol][iii]))
        ctr = 0
        for iant in range(nant):
         if iant+1!=27:
          ctr += 1
          pl.subplot(nn, mm, ctr)
          pl.xticks([])
          pl.plot(phs[iph][ipol,:,iant])#,'.', ms=0.2)
          pl.title(names[iant])
        pl.suptitle(title+' pol '+str(ipol), fontsize=20)
        pl.savefig(title.replace(' ','_')+'_pol'+str(ipol)+'.png', pad_inches=0.1, bbox_inches = 'tight')
  
    
################################################################################################
def plot_bandpass(caltable, calname=None, num=None, doplot=True):

    import matplotlib 
    from myrfimask import myrfimask as myrfi
    matplotlib.use("Agg")
    import pylab as pl
    pl.ion()

    if isinstance(caltable,str):
      data, antpos, names, err, snr, flag, u, v, index, dum = read_gains(caltable)
      calname = caltable
    else: 
      data, antpos, names, err, snr, flag, u, v, index, dum = caltable
    data, err, snr, flag = N.squeeze(data), N.squeeze(err), N.squeeze(snr), N.squeeze(flag)
    scans, nant, field, nchan, freq0, bw = dum
    print("Shape of data",data.shape)
    npol, nscan, nant, nchan = data.shape

    flagchans = []
    for i in range(len(myrfi)):
      x,y = myrfi[i]
      x,y = x*1e6, y*1e6
      ch1, ch2 = int((x-freq0)*nchan/bw), int((y-freq0)*nchan/bw)
      if ch1>0 and ch2>0 and ch1<nchan and ch2<nchan:
        flagchans.append([ch1,ch2])
    flagchans.append([0,int(0.04*nchan)])
    flagchans.append([int(0.92*nchan),nchan])

    if num == None:
      ff = glob.glob('15*')
      if len(ff)==0: snum = ''
      else: snum = ff[0].split('_')[0]
    else:
      snum = str(num)

    amp = N.abs(data)*flag
    nk = int(nchan/4096)  # 1k, 4k, 32k => 32, 8, 1
    ph = N.angle(data,1)*flag

    # amp is pol time ant chan
    if len(flagchans)>0:
      for fchan in flagchans:
        amp[:,:,:,fchan[0]:fchan[1]] = N.nan
        ph[:,:,:,fchan[0]:fchan[1]] = N.nan
    N.save('av_amp_'+calname, N.nanmean(N.nanmean(N.nanmean(amp,0),0),0))
    
    pl.figure(figsize=(10,8))
    avsp = N.nanmean(amp,1); avsp = N.nanmean(avsp,0); 
    for ia in range(avsp.shape[0]):
      pl.plot(avsp[ia], lw=1)
    avsp = N.nanmean(avsp,0)
    pl.plot(avsp,'k',lw=3)
    pl.suptitle(snum+' av bandpass amp')
    pl.savefig(snum+'_avbp_all.png')
    print("Wrote",snum+'_avbp_all.png')

    avsp = N.nanmean(amp,1); avsp = N.nanmean(avsp,0); avsp = N.nanmean(avsp,0) # av time then pol then ant
    stdsp = N.nanstd(amp,1); stdsp = N.nanmean(stdsp,0); stdsp = N.nanmean(stdsp,0)

    pl.figure(figsize=(10,8))
    pl.subplot(511); pl.plot(N.arange(len(avsp)), avsp)
    pl.title("av bandpass amp"); pl.xticks([])
    pl.suptitle(snum+" "+calname)
    xx = avsp-poly_filter(avsp,int(nchan/16))
    pl.subplot(512); pl.plot(N.arange(len(avsp)), xx); pl.xticks([])
    pl.title("Detrended "+str(int(nchan/16)))
    xav = N.ones(len(xx))*N.nan
    for i in range(55,len(xx)-55): 
      if N.sum(~N.isnan(xx[i-50:i+50]))>4: xav[i] = N.nanmean(xx[i-50:i+50])
    pl.subplot(513); pl.plot(xav); pl.xticks([])
    pl.title("Running mean "+str(100))
    pl.subplot(514); pl.plot(N.arange(len(stdsp)), stdsp); pl.xticks([])
    pl.title(" std bandpass amp")
    xx = stdsp#-poly_filter(stdsp,100)
    xav = N.ones(len(xx))*N.nan
    for i in range(55,len(xx)-55): 
      if N.sum(~N.isnan(xx[i-50:i+50]))>4: xav[i] = N.nanmean(xx[i-50:i+50])
    pl.subplot(515); pl.plot(xav); pl.title("Running mean of std "+str(100))
    pl.savefig(snum+'_'+calname+"_avbandpass.png")
    print("Wrote",snum+'_'+calname+"_avbandpass.png")
   
    #for ii in range(nant):
    #  for jj in range(2):
    #    gain = N.nanmedian(amp[jj,:,ii,:], 1)  # median gain vs scan
    #    #for ich in range(amp.shape[3]): amp[jj,:,ii,ich] -= gain

    c1, c2 = 0, nchan
    if doplot:
      pl.figure(figsize=(12,14))
      for iant in range(nant):
        pl.clf()
        avbp = N.nanmean(amp, 1)
  
        pl.subplot(421); libs.imshow(N.transpose(amp[0,:,iant,c1:c2])); pl.colorbar(); pl.title('Amp pol0')
        pl.subplot(423); libs.imshow(N.transpose(amp[1,:,iant,c1:c2])); pl.colorbar(); pl.title('Amp pol1')
        pl.subplot(425); pl.plot(N.nanmean(N.transpose(amp[0,:,iant,:])[int(ph.shape[3]*3/4)-15:int(ph.shape[3]*3/4)+15],0))
        pl.subplot(425); pl.plot(N.nanmean(N.transpose(amp[1,:,iant,:])[int(ph.shape[3]*3/4)-15:int(ph.shape[3]*3/4)+15],0))
        pl.xlabel('Scan')
        dum = N.nanmean(N.transpose(amp[0,:,iant,c1:c2]),1)
        dum = dum - median_filter(dum,100, mode='nearest')
        pl.subplot(427); pl.plot(N.arange(c1,c2), dum); pl.xlabel('Channel')
        dum = N.nanmean(N.transpose(amp[1,:,iant,c1:c2]),1)
        dum = dum - median_filter(dum,100, mode='nearest')
        pl.subplot(427); pl.plot(N.arange(c1,c2), dum); pl.xlabel('Channel')
  
        pl.subplot(422); libs.imshow(N.transpose(ph[0,:,iant,c1:c2])); pl.colorbar(); pl.title('Ph pol0')
        pl.subplot(424); libs.imshow(N.transpose(ph[1,:,iant,c1:c2])); pl.colorbar(); pl.title('Ph pol1')
        pl.subplot(426); xx = N.nanmean(N.transpose(ph[0,:,iant,:])[int(ph.shape[3]*3/4)-15:int(ph.shape[3]*3/4)+15],0); pl.plot(xx-N.nanmean(xx))
        pl.subplot(426); xx = N.nanmean(N.transpose(ph[1,:,iant,:])[int(ph.shape[3]*3/4)-35:int(ph.shape[3]*3/4)+35],0); pl.plot(xx-N.nanmean(xx))
        dum = N.nanmean(N.transpose(ph[:,:,iant,:]),1)
        pl.subplot(428); pl.plot(dum[:,0]-N.nanmean(dum[:,0])); pl.xlabel('Channel')
        pl.subplot(428); pl.plot(dum[:,1]-N.nanmean(dum[:,1])); pl.xlabel('Channel')
        pl.suptitle(snum+" "+calname+" Ant num "+str(iant+1))
        pl.savefig(snum+'_'+calname+"_bandpassplot_"+str(iant+1)+".png")
      print("Wrote",snum+'_'+calname+"_bandpassplot_*.png")
  
    
################################################################################################
    
def plot_gaincal(caltable, calname=None, num=None, fieldsep=False, t1=0, t2=0):

    from scipy.stats import pearsonr
    import matplotlib
    matplotlib.use("Agg")
    import pylab as pl
    pl.ion()

    if isinstance(caltable,str):
      data, antpos, names, err, snr, flag, u, v, index, dum = read_gains(caltable)
      calname = caltable
    else: 
      data, antpos, names, err, snr, flag, u, v, index, dum = caltable
    data, err, snr, flag = N.squeeze(data), N.squeeze(err), N.squeeze(snr), N.squeeze(flag)
    print('npol, ntime, nant = ', data.shape)
    npol, nscan, nant = data.shape

    if num == None:
      ff = glob.glob('15*')
      if len(ff)==0: snum = ''
      else: snum = ff[0].split('_')[0]
    else:
      snum = str(num)

    if caltable[:5]=='delay': amp = data*flag
    else: amp = N.abs(data)*flag  # pol time ant

    x, dum1, field, dum1, dum1, dum1 = dum
    if fieldsep==False: amp = [amp]
    else:
      nfd = len(N.unique(field)); amp1 = []
      for fd in range(nfd):
        inds = N.where(field==fd)[0]; 
        amp1.append(amp[:,inds])
      amp = amp1

    ampmed = []
    x = x[::nant]
    for jj in range(len(amp)): # nfields
      amp1 = N.copy(amp[jj])
      if fieldsep:
        inds = N.where(field==jj)[0]; x1 = x[inds]
      else:
        x1 = x
      inds = N.where(N.roll(x1,1)-x1 != 0)[0]
      start = N.copy(inds)
      end = list(inds[1:])
      end.append(len(x1))
      for ii in range(len(start)):
        med = N.nanmedian(amp[jj][:,start[ii]:end[ii],:], 1)   # pol time ant
        amp1[:,start[ii]:end[ii],:] = amp[jj][:,start[ii]:end[ii],:] - N.repeat(N.expand_dims(med,1),end[ii]-start[ii],1)
      ampmed.append(amp1)

    if caltable[:5]=='delay':
      ph = data*flag
    else:
      ph = N.angle(data,1)*flag
    med = N.nanmedian(ph, 1)
    ph = ph - N.repeat(N.expand_dims(med,1),nscan,1)
    ph = N.where(ph>180, ph-360, ph)
    ph = N.where(ph<-180, ph+360, ph)
    if fieldsep==False: ph = [ph]
    else:
      nfd = len(N.unique(field)); ph1 = []
      for fd in range(nfd):
        inds = N.where(field==fd)[0]; ph1.append(ph[:,inds])
      ph = ph1

    ph = N.asarray(ph)
    amp = N.asarray(amp)

    toplot = [amp, ph]
    titles = ['Amp', 'Phase_(-med)']

    nn, mm = libs.subplot(nant-1)
    for jj in range(len(amp)):
      for ii in range(2): # abs, angle
        pl.figure(figsize=(18,16))
        mn, mx = 1e7, -1e7
        for iii in range(len(amp)): 
          mn,mx = min(mn, N.nanmin(toplot[ii][iii])), max(mx, N.nanmax(toplot[ii][iii]))
        ctr = 0
        for iant in range(nant):
         if iant+1!=27:
          ctr += 1
          pl.subplot(nn, mm, ctr)
          if (ctr)%mm!=1: pl.yticks([])
          pl.xticks([])
          for ipol in range(2):
            pl.plot(toplot[ii][jj][ipol,:,iant])#,'.', ms=0.2)
            a1,a2,a3,a4=pl.axis()
            pl.axis([0, a2, mn, mx])
          pl.title(names[iant])
        #pl.suptitle(snum+' '+titles[ii]+' of '+calname+' versus time', fontsize=15)
        if fieldsep:
          titdum = snum+'_'+titles[ii]+'_'+calname+'_field'+str(jj)+'.png'
        else:
          titdum = snum+'_'+titles[ii]+'_'+calname+'.png'
        pl.savefig(titdum, pad_inches=0.1, bbox_inches = 'tight')
        print("Wrote",titdum)
          
    for jj in range(len(amp)):
      pl.figure(figsize=(20,12))
      pl.subplot(231); libs.imshow(N.transpose(amp[jj][0]))
      pl.xlabel('Antenna'); pl.ylabel('Scan num'); pl.title('Amp pol 0', fontsize=15); pl.colorbar()
      pl.subplot(234); libs.imshow(N.transpose(amp[jj][1]))
      pl.xlabel('Antenna'); pl.ylabel('Scan num'); pl.title('Amp pol 1', fontsize=15); pl.colorbar()
      pl.subplot(232); libs.imshow(N.transpose(ampmed[jj][0]))
      pl.xlabel('Antenna'); pl.ylabel('Scan num'); pl.title('Amp pol 0 (-med)', fontsize=15); pl.colorbar()
      pl.subplot(235); libs.imshow(N.transpose(ampmed[jj][1]))
      pl.xlabel('Antenna'); pl.ylabel('Scan num'); pl.title('Amp pol 1 (-med)', fontsize=15); pl.colorbar()
  
      pl.subplot(233); libs.imshow(N.transpose(ph[jj][0]))
      pl.xlabel('Antenna'); pl.ylabel('Scan num'); pl.title('Phase pol 0 (-med)', fontsize=15); pl.colorbar()
      pl.subplot(236); libs.imshow(N.transpose(ph[jj][1]))
      pl.xlabel('Antenna'); pl.ylabel('Scan num'); pl.title('Phase pol 1 (-med)', fontsize=15); pl.colorbar()
      pl.suptitle(snum+' '+calname+' field '+str(jj), fontsize=15)
      if fieldsep: titdum = snum+'_'+calname+'_field'+str(jj)+'_figs.png'
      else: titdum = snum+'_'+calname+'_figs.png'
      pl.savefig(titdum)
      print("Wrote",titdum)

      pl.figure(figsize=(12,12))
      for i in range(2):
        pl.subplot(221); pl.plot(N.nanstd(ampmed[jj][i],0)); pl.title('RMS of gainamp over time')
        pl.subplot(222); pl.plot(N.nanstd(ampmed[jj][i],1)); pl.title('RMS of gainamp over antenna')
        pl.subplot(223); pl.plot(N.nanstd(ph[jj][i],0)); pl.title('RMS of gain phase over time')
        pl.subplot(224); pl.plot(N.nanstd(ph[jj][i],1)); pl.title('RMS of gain phase over antenna')
      if fieldsep: titdum = snum+'_'+calname+'_field'+str(jj)+'_plots.png'
      else: titdum = snum+'_'+calname+'_plots.png'
      pl.suptitle(snum+' '+calname+'rms of amp and phase over time and ant')
      pl.savefig(titdum)
      print("Wrote", titdum)

    index0, index1, antnums, antind = index
    if caltable[:5]=='delay':
      ph = data*flag
    else:
      ph = N.squeeze(N.angle(data,1)*flag)

    na = ph.shape[2]
    nbl = int(na*(na-1)/2)
    cc, dd = [], []

    if t2==0: t2 = ph.shape[1]

    ph = ph[:,t1:t2,:]
    for ipol in range(2):
      cc1, dd1 = [], []; 
      for j in range(na):
        ph[ipol,:,j] -= N.nanmedian(ph[ipol,:,j])
        ph[ipol,:,j] = N.where(ph[ipol,:,j]>180, ph[ipol,:,j]-360, ph[ipol,:,j])
        ph[ipol,:,j] = N.where(ph[ipol,:,j]<-180, ph[ipol,:,j]+360, ph[ipol,:,j])
        ph[ipol,:,j] -= poly_filter(ph[ipol,:,j],ph.shape[1], order=3)

      uvdist = N.zeros(nbl)
      for ibl in range(nbl):
        uvdist[ibl] = N.sqrt(u[ibl]*u[ibl]+v[ibl]*v[ibl])
      indx = N.argsort(uvdist)[::-1]

      # plot 16 longest baselines
      pl.figure(figsize=(18,10))
      ctr = 1
      ddu, ddv = [], []; pairs=[]; uvbl = []
      for j in range(na):
        for k in range(na):
          if j>k: 
            xx, yy = ph[ipol,:,j],ph[ipol,:,k]
            ind = N.where((~N.isnan(ph[ipol,:,j]))*(~N.isnan(ph[ipol,:,k])))[0]
            if len(ind)>10: cc1.append(pearsonr(xx[ind], yy[ind])[0])
            else: cc1.append(N.nan)

            ddu.append(u[antind[j,k]])
            ddv.append(v[antind[j,k]])
            uvbl.append(uvdist[antind[j,k]])
            pairs.append([j,k])

            if (len(N.where((index0[indx[:16]]==j)*(index1[indx[:16]]==k))[0])==1) or \
               (len(N.where((index0[indx[:16]]==k)*(index1[indx[:16]]==j))[0])==1):
              pl.subplot(4,4,ctr); 
              if ctr<12: pl.xticks([])
              pl.plot(ph[ipol,:,j], lw=0.9); pl.plot(ph[ipol,:,k],lw=0.9)
              ctr += 1
              pl.title(names[j]+'&'+names[k], fontsize=15)
      pairs = N.asarray(pairs); uvbl = N.asarray(uvbl)
      cc.append(cc1); dd.append(dd1); cc1 = N.asarray(cc1)
      pl.suptitle(snum+' '+calname+' pol'+str(ipol)+' ant gain phase for highest uv', fontsize=15)
      pl.savefig(snum+'_'+calname+'_pol'+str(ipol)+'_highuv.png')
      print("Wrote",snum+'_'+calname+'_pol'+str(ipol)+'_highuv.png')

      pl.figure(figsize=(12,10))
      pl.scatter(ddu, ddv, c=cc1, cmap='jet')
      pl.colorbar()
      pl.savefig(snum+'_'+calname+'_uv_cc_pol'+str(ipol)+'.png')
      print("Wrote",snum+'_'+calname+'_uv_cc_pol'+str(ipol)+'.png')

      pl.figure(figsize=(12,10))
      pl.scatter(antpos[:,0], antpos[:,1])
      for ii in range(len(cc1)):
       if uvbl[ii]>1000:
        j, k = pairs[ii]
        if cc1[ii]>0.5: pl.plot([antpos[j,0], antpos[k,0]], [antpos[j,1], antpos[k,1]], '-b')
        if cc1[ii]<-0.5: pl.plot([antpos[j,0], antpos[k,0]], [antpos[j,1], antpos[k,1]], '-r')
      pl.savefig(snum+'_'+calname+'_xy_cc_pol'+str(ipol)+'.png')
      print("Wrote",snum+'_'+calname+'_xy_cc_pol'+str(ipol)+'.png')
      # plot random bls in each uv range
      ranges = [[0,1000],[1500,2500],[3000,4000], [5500,8000]]  # random pairs in this uv range and cross corr > 0.3 in abs
      pnum = 5
      for ii in range(len(ranges)):
        pl.figure(figsize=(12,10))
        ctr = 1
        a1, a2 = N.random.randint(0,na,(2,pnum*pnum))
        distind = N.where((ranges[ii][0]<uvbl)*(uvbl<ranges[ii][1])*(N.abs(cc1)>0.5))[0]
        if len(distind)>0:
          myrandind = distind[N.random.randint(0,len(distind),pnum*pnum)]
          a1 = pairs[myrandind][:,0]; a2 = pairs[myrandind][:,1]
          for jj,j in enumerate(a1):
             k = a2[jj]
             pl.subplot(pnum,pnum,jj+1); 
             pl.plot(ph[ipol,:,j],lw=0.9); pl.plot(ph[ipol,:,k],lw=0.9)
             #df = [antpos[j,0]-antpos[k,0], antpos[j,1]-antpos[k,1]]
             #dist = N.sqrt(df[0]*df[0]+df[1]*df[1])
             xx, yy = ph[ipol,:,j],ph[ipol,:,k]
             ind = N.where((~N.isnan(ph[ipol,:,j]))*(~N.isnan(ph[ipol,:,k])))[0]
             if len(ind)>3: dum2 = pearsonr(xx[ind], yy[ind])[0]
             if jj+1<pnum*(pnum-1)+1: pl.xticks([])
             nam1, nam2 = names[j][2:], names[k][2:]
             pl.title("%s %d%s %s%.2f"%(nam1+'-'+nam2,int(uvdist[antind[j,k]]),'m',' r=',dum2))
          pl.suptitle(snum+' '+calname+' pol'+str(ipol)+' ant gain phase for rand bl in '+str(ranges[ii]))
          pl.savefig(snum+'_'+calname+'_pol'+str(ipol)+'_randomuv_'+str(ii)+'.png')
          print("Wrote",snum+'_'+calname+'_pol'+str(ipol)+'_randomuv_'+str(ii)+'.png')
    cc, dd = N.asarray(cc), N.asarray(dd)

    pl.figure(figsize=(10,6))
    pl.plot(uvbl, cc[0], '.', ms=5)
    pl.plot(uvbl, cc[1], '.', ms=5,alpha=0.5)
    a1,a2,a3,a4 = pl.axis(); #pl.axis([a1,a2,-0.5,1.0])
    pl.plot([a1,a2],[0,0], 'k-')
    pl.xlabel('uv dist (m)'); pl.ylabel('Corr coeff')
    pl.title(snum+' '+calname+' coeff corr bet gain phases of every ant pair vs uv dist')
    pl.savefig(snum+'_'+calname+'_phasecc.png')
    print("Wrote",snum+'_'+calname+'_phasecc.png')

  
################################################################################################

def gainphasephase(caltable):
    import scipy.stats as ss
    import katdal
    import pylab as pl
    pl.ion()

    data, err, snr, flag, u, v, dum = read_gains(caltable)
    data, err, snr, flag = N.squeeze(data), N.squeeze(err), N.squeeze(snr), N.squeeze(flag)
    print('npol, nscan, nant', data.shape, err.shape, snr.shape, flag.shape)
    npol, nscan, nant = data.shape 

    ff = glob.glob('15*')
    if len(ff)==0: snum = ''
    else: snum = ff[0].split('_')[0]

    amp = N.abs(data)*flag
    ph = N.angle(data,1)*flag
    med = N.nanmedian(ph, 1)
    ph = ph - N.repeat(N.expand_dims(med,1),nscan,1)
    ph = N.where(ph>180, ph-360, ph)
    ph = N.where(ph<-180, ph+360, ph)

    pl.figure()
    phph = N.ones((nant,2))
    for iant in range(nant):
      for ipol in range(2):
        yy = ph[ipol,:,iant]
        yy = ph[ipol,900:1600,iant]
        yy -= N.nanmean(yy)
        xx = N.arange(len(yy))
        dum = ss.linregress(xx,yy)
        mm, cc = dum[:2]
        yy = yy - mm*xx-cc
        ff = fft1(yy)
        fnu, famp, fph, dum = ff
        roll = famp-N.roll(famp,1)
        print(xx, yy, fnu, famp)
        print()
        if iant>nant-23:
            pl.subplot(6,4,iant-nant+23)
            pl.plot(famp[300:int(len(famp)/2)])

        posn = N.argmax(famp[:int(len(famp)/2)-5])
        dum = fph[posn]
        if dum<0: dum += 180
        phph[iant,ipol] = dum
          
    pl.savefig('fft.png')

    fn = glob.glob('15*rdb')[0]
    f = katdal.open(fn)
    x0, y0 = f.ants[0].position_ecef[:2]
    pl.figure(figsize=(16,8))
    for ipol in range(2):
      pl.subplot(1,2,ipol+1)
      x = [f.ants[iant].position_ecef[0]-x0 for iant in range(1,nant)]
      y = [f.ants[iant].position_ecef[1]-y0 for iant in range(1,nant)]
      pl.scatter(x, y, c=phph[1:,ipol], cmap='jet')
    pl.suptitle(snum+' phase of the time-ripples in gain phases of each antenna for each pol')
    pl.savefig(snum+'_phase_ripple_gainphase_antposn.png')

################################################################################################




      

 
        










